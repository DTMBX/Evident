# BarberX Legal Tech Platform
# Professional BWC Forensic Analysis System
# Copyright (c) 2026 BarberX Legal Technologies

import hashlib
import json
import logging
import os
import sys
import threading
import uuid
from datetime import datetime, timedelta
from logging.handlers import RotatingFileHandler
from pathlib import Path

import flask
from flask import (Flask, flash, jsonify, redirect, render_template, request,
                   send_file, session, url_for)
from flask_compress import Compress
from flask_cors import CORS
from flask_login import (LoginManager, current_user, login_required,
                         login_user, logout_user)
from flask_wtf.csrf import CSRFProtect
from werkzeug.utils import secure_filename

from free_tier_data_retention import DataRetentionManager, get_user_data_status
# FREE Tier Functionality
from free_tier_demo_cases import (get_demo_case_by_id, get_demo_cases,
                                  is_demo_case)
from free_tier_educational_resources import (CATEGORIES,
                                             get_all_educational_resources,
                                             get_resource_by_id)
from free_tier_upload_manager import (OneTimeUploadManager,
                                      free_tier_upload_route_decorator)
from free_tier_watermark import WatermarkService

# Security utilities
try:
    from utils.security import ErrorSanitizer

    UTILS_AVAILABLE = True
except ImportError as e:
    UTILS_AVAILABLE = False
    print(f"[!] Utils not available: {e}")

    # Fallback ErrorSanitizer
    class ErrorSanitizer:
        @staticmethod
        def create_error_ticket():
            return "ERR-" + str(uuid.uuid4())[:8]

        @staticmethod
        def sanitize_error(error):
            return "An error occurred. Please contact support."


# Enhanced authentication imports (CRITICAL: TierLevel used in decorators at module level)
try:
    from auth_routes import auth_bp

    ENHANCED_AUTH_AVAILABLE = True
except ImportError as e:
    ENHANCED_AUTH_AVAILABLE = False
    print(f"[!] Enhanced auth not available: {e}")

# Import TierLevel and User at module level (required for decorators)
from models_auth import ApiKey as APIKey
from models_auth import TierLevel, User

# Tier gating system
try:
    from tier_gating import check_usage_limit, require_tier

    TIER_GATING_AVAILABLE = True
except ImportError as e:
    TIER_GATING_AVAILABLE = False
    print(f"[!] Tier gating not available: {e}")

    # Fallback decorators that do nothing
    def require_tier(tier):
        def decorator(f):
            return f

        return decorator

    def check_usage_limit(field, increment=0, hours=None):
        def decorator(f):
            return f

        return decorator


# UX enhancement utilities
try:
    from ux_helpers import register_ux_filters

    UX_HELPERS_AVAILABLE = True
except ImportError as e:
    UX_HELPERS_AVAILABLE = False
    print(f"[!] UX helpers not available: {e}")

# Phase 1 Premium Features
try:
    from whisper_transcription import WhisperTranscriptionService

    WHISPER_AVAILABLE = True
except ImportError as e:
    WHISPER_AVAILABLE = False
    print(f"[!] Whisper transcription not available: {e}")

try:
    from ocr_service import OCRService

    OCR_AVAILABLE = True
except ImportError as e:
    OCR_AVAILABLE = False
    print(f"[!] OCR service not available: {e}")

try:
    from two_factor_auth import TwoFactorAuthService

    TWO_FACTOR_AVAILABLE = True
except ImportError as e:
    TWO_FACTOR_AVAILABLE = False
    print(f"[!] 2FA service not available: {e}")

try:
    from stripe_payment_service import SUBSCRIPTION_PLANS, StripePaymentService

    STRIPE_AVAILABLE = True
except ImportError as e:
    STRIPE_AVAILABLE = False
    print(f"[!] Stripe payment service not available: {e}")

# Import our BWC analyzer (optional - only needed for actual analysis)
try:
    pass

    BWC_ANALYZER_AVAILABLE = True
except ImportError:
    BWC_ANALYZER_AVAILABLE = False
    app_logger = logging.getLogger(__name__)
    app_logger.warning("BWC Forensic Analyzer not available - AI dependencies not installed")

# Backend Optimization Components
try:
    from backend_integration import (error_response, event_bus,
                                     service_registry, success_response)
    from config_manager import ConfigManager, DatabaseOptimizer
    from unified_evidence_service import (EvidenceReportGenerator,
                                          UnifiedEvidenceProcessor)

    BACKEND_OPTIMIZATION_AVAILABLE = True
    print("[OK] Backend optimization components loaded")
except ImportError as e:
    BACKEND_OPTIMIZATION_AVAILABLE = False
    print(f"[!] Backend optimization not available: {e}")

    # Fallback response helpers
    def success_response(message, data=None):
        return {"success": True, "message": message, "data": data or {}}

    def error_response(message, error_code=None, details=None):
        return {"success": False, "message": message, "error_code": error_code, "details": details}


# Initialize Flask app
app = Flask(__name__)

# Initialize logger for the application
logger = logging.getLogger(__name__)

# Enable compression for all responses
compress = Compress()

# Initialize CSRF protection
csrf = CSRFProtect()

# Load environment variables
from dotenv import load_dotenv

load_dotenv()

# Initialize configuration manager (if available)
if BACKEND_OPTIMIZATION_AVAILABLE:
    config_mgr = ConfigManager()
    # Apply SQLAlchemy configuration
    app.config.update(config_mgr.get_sqlalchemy_config())
    app.config["SECRET_KEY"] = config_mgr.config.secret_key
    app.config["UPLOAD_FOLDER"] = Path(config_mgr.config.upload_folder)
    app.config["MAX_CONTENT_LENGTH"] = config_mgr.config.max_upload_size
    print(f"[OK] Configuration manager initialized - {config_mgr.config.environment} environment")
else:
    # Fallback to manual configuration
    # CRITICAL: SECRET_KEY must be set in environment - no hardcoded fallback
    secret_key = os.getenv("SECRET_KEY")
    if not secret_key:
        if app.config.get("TESTING"):
            # Only allow auto-generation in testing
            import secrets

            secret_key = secrets.token_hex(32)
            app.logger.warning("Generated temporary SECRET_KEY for testing")
        else:
            raise RuntimeError(
                "SECRET_KEY environment variable is required for security. "
                "Generate one with: python -c 'import secrets; print(secrets.token_hex(32))'"
            )
    app.config["SECRET_KEY"] = secret_key

    # Use absolute path for database
    basedir = os.path.abspath(os.path.dirname(__file__))

    # Database configuration - PostgreSQL for production, SQLite for development
    database_url = os.getenv("DATABASE_URL")
    if database_url:
        # Fix for Heroku/Render postgres URL
        if database_url.startswith("postgres://"):
            database_url = database_url.replace("postgres://", "postgresql://", 1)
        app.config["SQLALCHEMY_DATABASE_URI"] = database_url
        print("[OK] Using PostgreSQL database for production")
    else:
        # Local development with SQLite
        app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///barberx_FRESH.db"
        print("[OK] Using SQLite database for development")

    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
    app.config["MAX_CONTENT_LENGTH"] = int(
        os.getenv("MAX_CONTENT_LENGTH", 20 * 1024 * 1024 * 1024)
    )  # 20GB max to support Enterprise tier

    # [SECURITY FIX] Cookie security flags to prevent session hijacking and CSRF
    # See: https://owasp.org/www-community/controls/SecureCookieAttribute
    is_production = os.getenv("FLASK_ENV") == "production" or os.getenv("FORCE_HTTPS") == "true"
    app.config["SESSION_COOKIE_SECURE"] = is_production  # Only send over HTTPS in production
    app.config["SESSION_COOKIE_HTTPONLY"] = True  # Prevent JavaScript access (already set)
    app.config["SESSION_COOKIE_SAMESITE"] = "Lax"  # Prevent CSRF attacks while allowing normal nav
    app.config["UPLOAD_FOLDER"] = Path("./uploads/bwc_videos")

app.config["ANALYSIS_FOLDER"] = Path("./bwc_analysis")
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)

# Initialize AI Pipeline Orchestrator with configuration
try:
    from src.ai.pipeline import get_orchestrator

    pipeline_config = {
        "storage_root": "./uploads/pdfs/originals",
        "manifest_root": "./manifest",
        "db_path": "instance/barberx_legal.db",
        "ocr_threshold": 50,
        "enable_citation_tracking": True,
        "max_passage_length": 1500,
        "retrieval_top_k": 10,
    }

    # Initialize the singleton orchestrator
    orchestrator = get_orchestrator(pipeline_config)
    print("[OK] AI Pipeline orchestrator initialized with citation tracking")
except ImportError as e:
    print(f"[WARN] AI Pipeline not available: {e}")

# Stripe Configuration (from environment)
app.config["STRIPE_PRICING_TABLE_ID"] = os.getenv("STRIPE_PRICING_TABLE_ID")
app.config["STRIPE_PUBLISHABLE_KEY"] = os.getenv("STRIPE_PUBLISHABLE_KEY")
app.config["STRIPE_SECRET_KEY"] = os.getenv("STRIPE_SECRET_KEY")
app.config["STRIPE_WEBHOOK_SECRET"] = os.getenv("STRIPE_WEBHOOK_SECRET")

# CORS configuration for production (includes mobile/desktop clients)
cors_origins = os.getenv(
    "CORS_ORIGINS",
    "https://barberx.info,https://www.barberx.info,http://localhost:5000,http://127.0.0.1:5000,tauri://localhost,capacitor://localhost",
)
CORS_ORIGINS_LIST = [origin.strip() for origin in cors_origins.split(",")]

# Create directories
app.config["UPLOAD_FOLDER"].mkdir(parents=True, exist_ok=True)
app.config["ANALYSIS_FOLDER"].mkdir(parents=True, exist_ok=True)

# File Upload Security Configuration
ALLOWED_VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".flv", ".wmv", ".m4v"}
ALLOWED_AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".aac", ".ogg", ".wma", ".flac"}
ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".rtf"}
ALLOWED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".svg"}

ALL_ALLOWED_EXTENSIONS = (
    ALLOWED_VIDEO_EXTENSIONS
    | ALLOWED_AUDIO_EXTENSIONS
    | ALLOWED_DOCUMENT_EXTENSIONS
    | ALLOWED_IMAGE_EXTENSIONS
)

ALLOWED_MIME_TYPES = {
    # Video
    "video/mp4",
    "video/quicktime",
    "video/x-msvideo",
    "video/x-matroska",
    "video/webm",
    "video/x-flv",
    "video/x-ms-wmv",
    # Audio
    "audio/mpeg",
    "audio/wav",
    "audio/mp4",
    "audio/aac",
    "audio/ogg",
    "audio/x-ms-wma",
    "audio/flac",
    # Documents
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "application/rtf",
    # Images
    "image/jpeg",
    "image/png",
    "image/gif",
    "image/bmp",
    "image/tiff",
    "image/svg+xml",
}

# File size limits (can be overridden per tier)
MAX_FILE_SIZE_FREE = 500 * 1024 * 1024  # 500MB for free tier
MAX_FILE_SIZE_PRO = 2 * 1024 * 1024 * 1024  # 2GB for pro tier
MAX_FILE_SIZE_PREMIUM = 5 * 1024 * 1024 * 1024  # 5GB for premium tier


# Initialize extensions (defer db binding)
from models_auth import db

db.init_app(app)
compress.init_app(app)
csrf.init_app(app)

# Exempt Stripe webhook from CSRF (it uses signature verification instead)
csrf.exempt("stripe_payments.webhook")
csrf.exempt("stripe.stripe_webhook")

# Exempt REST API endpoints from CSRF (they use JWT/session authentication)
csrf.exempt("auth_api")
csrf.exempt("upload_api")
csrf.exempt("analysis_api")
csrf.exempt("user_api")
csrf.exempt("stripe_api")
csrf.exempt("admin_api")
csrf.exempt("evidence_api")

CORS(
    app,
    origins=CORS_ORIGINS_LIST,
    supports_credentials=True,
    allow_headers=["Content-Type", "Authorization"],
    methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
)
login_manager = LoginManager(app)
login_manager.login_view = "auth.login"  # Updated to use auth blueprint

# Configure request timeout
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = 31536000  # 1 year for static files
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)


# Security Headers Middleware (CSP, HSTS, etc.)
@app.after_request
def add_security_headers(response):
    """Add comprehensive security headers to all responses"""

    # Content Security Policy - Strict but allows necessary resources
    csp_policy = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval' https://js.stripe.com https://cdn.amplitude.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "font-src 'self' https://fonts.gstatic.com data:; "
        "img-src 'self' data: https: blob:; "
        "media-src 'self' blob:; "
        "connect-src 'self' https://api.stripe.com https://api.amplitude.com https://api.openai.com; "
        "frame-src 'self' https://js.stripe.com https://hooks.stripe.com; "
        "frame-ancestors 'none'; "
        "base-uri 'self'; "
        "form-action 'self'; "
        "upgrade-insecure-requests"
    )
    response.headers["Content-Security-Policy"] = csp_policy

    # HTTP Strict Transport Security - Force HTTPS for 1 year
    if request.is_secure or os.getenv("FORCE_HTTPS") == "true":
        response.headers["Strict-Transport-Security"] = (
            "max-age=31536000; includeSubDomains; preload"
        )

    # Prevent clickjacking attacks
    response.headers["X-Frame-Options"] = "DENY"

    # Prevent MIME type sniffing
    response.headers["X-Content-Type-Options"] = "nosniff"

    # Enable XSS protection in older browsers
    response.headers["X-XSS-Protection"] = "1; mode=block"

    # Referrer policy - don't leak information to external sites
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"

    # Permissions policy - restrict browser features
    response.headers["Permissions-Policy"] = (
        "geolocation=(), "
        "microphone=(), "
        "camera=(self), "
        'payment=(self "https://js.stripe.com"), '
        "usb=(), "
        "magnetometer=(), "
        "gyroscope=(), "
        "speaker=(self)"
    )

    # Cross-Origin policies
    response.headers["Cross-Origin-Embedder-Policy"] = "require-corp"
    response.headers["Cross-Origin-Opener-Policy"] = "same-origin"
    response.headers["Cross-Origin-Resource-Policy"] = "same-origin"

    return response


# Register enhanced authentication blueprint
if ENHANCED_AUTH_AVAILABLE:
    app.register_blueprint(auth_bp, url_prefix="/auth")
    print("[OK] Enhanced auth routes registered at /auth/*")

# Register batch upload handler
try:
    from batch_upload_handler import batch_upload_bp

    app.register_blueprint(batch_upload_bp)
    print("[OK] Unified batch upload registered at /api/upload/batch")
except ImportError as e:
    print(f"[WARN] Batch upload handler not available: {e}")

# Register Stripe payments blueprint
try:
    from stripe_payments import payments_bp

    app.register_blueprint(payments_bp)
    print("[OK] Stripe payments registered at /payments/*")
except ImportError as e:
    print(f"[WARN] Stripe payments not available: {e}")

# Register Stripe subscription service (webhooks)
try:
    from stripe_subscription_service import stripe_bp

    app.register_blueprint(stripe_bp)
    csrf.exempt(stripe_bp)  # Exempt entire blueprint - uses Stripe signature verification
    print("[OK] Stripe subscription service registered at /api/stripe/*")
except ImportError as e:
    print(f"[WARN] Stripe subscription service not available: {e}")

# Register ChatGPT integration blueprint
try:
    from api.chatgpt import chatgpt_bp

    app.register_blueprint(chatgpt_bp)
    csrf.exempt(chatgpt_bp)  # API endpoints use token auth, not CSRF
    print("[OK] ChatGPT integration registered at /api/v1/chat/*, /api/v1/projects/*")
except ImportError as e:
    print(f"[WARN] ChatGPT integration not available: {e}")

# Register Document Optimizer blueprint
try:
    from api.document_optimizer import bp as doc_optimizer_bp

    app.register_blueprint(doc_optimizer_bp)
    csrf.exempt(doc_optimizer_bp)  # API endpoints use token auth
    print("[OK] Document Optimizer registered at /api/document-optimizer/*")
except ImportError as e:
    print(f"[WARN] Document Optimizer not available: {e}")

# Register Legal Reference Library blueprint
try:
    from api.legal_library import bp as legal_library_bp

    app.register_blueprint(legal_library_bp)
    print("[OK] Legal Library registered at /api/legal-library/*")
except ImportError as e:
    print(f"[WARN] Legal Library not available: {e}")

# Register Enhanced Chat Assistant blueprint
try:
    from api.enhanced_chat import chat_bp

    app.register_blueprint(chat_bp)
    # Exempt after registration
    csrf.exempt(chat_bp)
    print("[OK] Enhanced Chat registered at /api/chat/*")
except ImportError as e:
    print(f"[WARN] Enhanced Chat not available: {e}")

# Register REST API blueprints for cross-platform clients
try:
    from api import register_api_blueprints

    register_api_blueprints(app)
except ImportError as e:
    print(f"[WARN] REST API blueprints not available: {e}")

# Register API Usage Metering blueprint
try:
    from api.metering import metering_bp

    app.register_blueprint(metering_bp)
    csrf.exempt(metering_bp)
    print("[OK] API Usage Metering registered at /api/v1/metering/*")
except ImportError as e:
    print(f"[WARN] API Usage Metering not available: {e}")

# Register Legal Chatbot API blueprint
try:
    from api.legal_chatbot import legal_chatbot_bp

    app.register_blueprint(legal_chatbot_bp)
    csrf.exempt(legal_chatbot_bp)
    print("[OK] Legal Chatbot API registered at /api/v1/chatbot/*")
except ImportError as e:
    print(f"[WARN] Legal Chatbot API not available: {e}")

# Register UX helper filters and context processors
if UX_HELPERS_AVAILABLE:
    register_ux_filters(app)
    print("[OK] UX enhancement filters registered")

# Configure logging
if not app.debug:
    if not os.path.exists("logs"):
        os.mkdir("logs")
    file_handler = RotatingFileHandler("logs/barberx.log", maxBytes=10240000, backupCount=10)
    file_handler.setFormatter(
        logging.Formatter("%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]")
    )
    file_handler.setLevel(logging.INFO)
    app.logger.addHandler(file_handler)
    app.logger.setLevel(logging.INFO)
    app.logger.info("BarberX Legal Tech startup")

# Initialize backend services (after database init)
evidence_processor = None
report_generator = None


def initialize_backend_services():
    """Initialize backend optimization services"""
    global evidence_processor, report_generator

    if not BACKEND_OPTIMIZATION_AVAILABLE:
        return

    try:
        # Create database indexes for performance
        optimizer = DatabaseOptimizer(db)
        optimizer.create_indexes()
        app.logger.info("[OK] Database indexes created/verified")

        # Initialize evidence processor
        evidence_processor = UnifiedEvidenceProcessor()
        report_generator = EvidenceReportGenerator()
        app.logger.info("[OK] Evidence processor initialized")

        # Subscribe to events (optional)
        def on_evidence_processed(event):
            app.logger.info(f"Evidence {event.data.get('evidence_id')} processed")

        event_bus.subscribe("evidence.processed", on_evidence_processed)
        event_bus.subscribe(
            "evidence.processing_failed",
            lambda e: app.logger.error(f"Processing failed: {e.data}"),
        )

        app.logger.info("[OK] Backend optimization services initialized")

    except Exception as e:
        app.logger.error(f"Failed to initialize backend services: {e}")


# Global analyzer instance (lazy load)
bwc_analyzer_instance = None  # noqa: F811 - Flask route shadow
analysis_tasks = {}  # Track background analysis tasks

# ========================================
# DATABASE MODELS
# ========================================
# Note: User and ApiKey imported from models_auth.py
# Analysis, AppSettings, PDFUpload, AuditLog defined here


class Analysis(db.Model):
    """Analysis record for BWC video processing"""

    __tablename__ = "analyses"
    __table_args__ = {"extend_existing": True}

    id = db.Column(db.String(32), primary_key=True)  # UUID
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)

    # File information
    filename = db.Column(db.String(255), nullable=False)
    file_hash = db.Column(db.String(64), nullable=False, index=True)
    file_size = db.Column(db.BigInteger, nullable=False)
    file_path = db.Column(db.String(500), nullable=False)

    # Case metadata
    case_number = db.Column(db.String(100), index=True)
    evidence_number = db.Column(db.String(100))
    acquired_by = db.Column(db.String(100))
    source = db.Column(db.String(200))
    known_officers = db.Column(db.JSON)

    # Analysis status
    status = db.Column(db.String(20), default="uploaded")
    progress = db.Column(db.Integer, default=0)
    current_step = db.Column(db.String(100))
    error_message = db.Column(db.Text)

    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)
    started_at = db.Column(db.DateTime)
    completed_at = db.Column(db.DateTime)

    # Results summary
    duration = db.Column(db.Float)
    total_speakers = db.Column(db.Integer)
    total_segments = db.Column(db.Integer)
    total_discrepancies = db.Column(db.Integer)
    critical_discrepancies = db.Column(db.Integer)

    # Report paths
    report_json_path = db.Column(db.String(500))
    report_txt_path = db.Column(db.String(500))
    report_md_path = db.Column(db.String(500))

    # Sharing and collaboration
    is_shared = db.Column(db.Boolean, default=False)
    share_token = db.Column(db.String(64), unique=True)

    # Tags for organization
    tags = db.Column(db.JSON)
    notes = db.Column(db.Text)

    def generate_id(self):
        """Generate unique analysis ID"""
        self.id = uuid.uuid4().hex

    def to_dict(self, include_results=False):
        data = {
            "id": self.id,
            "filename": self.filename,
            "file_hash": self.file_hash,
            "file_size": self.file_size,
            "case_number": self.case_number,
            "evidence_number": self.evidence_number,
            "status": self.status,
            "progress": self.progress,
            "created_at": self.created_at.isoformat(),
            "tags": self.tags or [],
        }

        if include_results and self.status == "completed":
            data.update(
                {
                    "duration": self.duration,
                    "total_speakers": self.total_speakers,
                    "total_segments": self.total_segments,
                    "total_discrepancies": self.total_discrepancies,
                    "critical_discrepancies": self.critical_discrepancies,
                    "completed_at": self.completed_at.isoformat() if self.completed_at else None,
                }
            )

        return data


class AppSettings(db.Model):
    """Application settings and configuration"""

    __tablename__ = "app_settings"
    __table_args__ = {"extend_existing": True}

    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(db.String(100), unique=True, nullable=False, index=True)
    value = db.Column(db.Text)
    value_type = db.Column(db.String(20), default="string")  # string, int, float, bool, json
    category = db.Column(
        db.String(50), default="general"
    )  # general, security, features, limits, email, branding
    description = db.Column(db.String(500))
    is_editable = db.Column(db.Boolean, default=True)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    updated_by = db.Column(db.Integer, db.ForeignKey("users.id"))

    @staticmethod
    def get(key, default=None):
        """Get setting value by key"""
        setting = AppSettings.query.filter_by(key=key).first()
        if not setting:
            return default

        # Convert based on type
        if setting.value_type == "bool":
            return setting.value.lower() in ("true", "1", "yes")
        elif setting.value_type == "int":
            return int(setting.value)
        elif setting.value_type == "float":
            return float(setting.value)
        elif setting.value_type == "json":
            import json

            return json.loads(setting.value)
        return setting.value

    @staticmethod
    def set(key, value, value_type="string", category="general", description=""):
        """Set or update setting value"""
        setting = AppSettings.query.filter_by(key=key).first()

        # Convert value to string for storage
        if value_type == "json":
            import json

            value_str = json.dumps(value)
        else:
            value_str = str(value)

        if setting:
            setting.value = value_str
            setting.value_type = value_type
            setting.updated_at = datetime.utcnow()
            if current_user.is_authenticated:
                setting.updated_by = current_user.id
        else:
            setting = AppSettings(
                key=key,
                value=value_str,
                value_type=value_type,
                category=category,
                description=description,
                updated_by=current_user.id if current_user.is_authenticated else None,
            )
            db.session.add(setting)

        db.session.commit()
        return setting


def get_site_settings():
    """Return site-wide UI controls with defaults."""
    defaults = {
        "ui.theme": "system",
        "ui.brand_color": "#c41e3a",
        "ui.accent_color": "#1e40af",
        "site.maintenance_mode": False,
        "site.maintenance_message": "We are performing maintenance. Please check back shortly.",
        "site.banner_enabled": False,
        "site.banner_message": "",
        "site.allow_signup": True,
        "site.footer_notice": "",
        "site.chat_enabled": True,
    }

    settings = {}
    for key, default in defaults.items():
        settings[key] = AppSettings.get(key, default)
    return settings


@app.context_processor
def inject_site_settings():
    """Expose site settings to all Jinja templates."""
    return {"site_settings": get_site_settings()}


# Template Filters
@app.template_filter("format_number")
def format_number_filter(value):
    """Format numbers with thousands separator"""
    try:
        return f"{int(value):,}"
    except (ValueError, TypeError):
        return value


@app.before_request
def enforce_site_controls():
    """Enforce maintenance mode and feature flags across the site."""
    settings = get_site_settings()
    path = request.path or ""

    allow_paths = (
        "/static",
        "/assets",
        "/favicon.ico",
        "/auth/login",
        "/auth/logout",
        "/admin",
    )

    if settings.get("site.maintenance_mode"):
        is_admin = current_user.is_authenticated and getattr(current_user, "role", "") == "admin"
        if not path.startswith(allow_paths) and not is_admin:
            return render_template("maintenance.html"), 503

    if not settings.get("site.allow_signup") and (
        path.startswith("/auth/signup")
        or path.startswith("/register")
        or path.startswith("/signup")
    ):
        flash("Signups are temporarily disabled.", "warning")
        return redirect("/auth/login")

    if not settings.get("site.chat_enabled"):
        if path.startswith("/chat"):
            flash("Chat is temporarily disabled.", "warning")
            return redirect(url_for("dashboard"))
        if path.startswith("/api/chat"):
            return jsonify({"error": "Chat is temporarily disabled."}), 503


class PDFUpload(db.Model):
    """Model for tracking uploaded PDF files"""

    __tablename__ = "pdf_uploads"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True, index=True)

    # File information
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    file_hash = db.Column(db.String(64), unique=True, nullable=False, index=True)
    mime_type = db.Column(db.String(100), default="application/pdf")

    # Metadata
    case_number = db.Column(db.String(100), index=True)
    document_type = db.Column(db.String(100))  # brief, motion, order, filing, etc.
    tags = db.Column(db.JSON)
    description = db.Column(db.Text)

    # Status and processing
    status = db.Column(db.String(20), default="uploaded")  # uploaded, processing, processed, error
    page_count = db.Column(db.Integer)
    extracted_text = db.Column(db.Text)

    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)
    processed_at = db.Column(db.DateTime)

    # Access control
    is_public = db.Column(db.Boolean, default=False)
    share_token = db.Column(db.String(64), unique=True)

    def generate_hash(self, file_path):
        """Generate SHA-256 hash of file"""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        self.file_hash = sha256.hexdigest()

    def to_dict(self):
        return {
            "id": self.id,
            "filename": self.filename,
            "original_filename": self.original_filename,
            "file_size": self.file_size,
            "file_hash": self.file_hash,
            "case_number": self.case_number,
            "document_type": self.document_type,
            "tags": self.tags or [],
            "description": self.description,
            "status": self.status,
            "page_count": self.page_count,
            "created_at": self.created_at.isoformat(),
            "processed_at": self.processed_at.isoformat() if self.processed_at else None,
        }


class AuditLog(db.Model):
    """Audit log for compliance and security"""

    __tablename__ = "audit_logs"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), index=True)

    action = db.Column(db.String(50), nullable=False, index=True)
    resource_type = db.Column(db.String(50))
    resource_id = db.Column(db.String(100))
    details = db.Column(db.JSON)

    ip_address = db.Column(db.String(45))
    user_agent = db.Column(db.String(500))

    created_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)

    @staticmethod
    def log(action, resource_type=None, resource_id=None, details=None):
        """Create audit log entry"""
        log_entry = AuditLog(
            user_id=current_user.id if current_user.is_authenticated else None,
            action=action,
            resource_type=resource_type,
            resource_id=resource_id,
            details=details,
            ip_address=request.remote_addr,
            user_agent=request.headers.get("User-Agent"),
        )
        db.session.add(log_entry)
        db.session.commit()


# ========================================
# AUTHENTICATION
# ========================================


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def api_key_required(f):
    """Decorator to require API key authentication"""
    from functools import wraps

    @wraps(f)
    def decorated_function(*args, **kwargs):
        api_key = request.headers.get("X-API-Key")

        if not api_key:
            return jsonify({"error": "API key required"}), 401

        key_obj = APIKey.query.filter_by(key=api_key, is_active=True).first()

        if not key_obj:
            return jsonify({"error": "Invalid API key"}), 401

        # Update last used
        key_obj.last_used_at = datetime.utcnow()
        db.session.commit()

        # Set current user
        request.current_api_user = key_obj.user

        return f(*args, **kwargs)

    return decorated_function


# ========================================
# HELPER FUNCTIONS
# ========================================


def validate_upload_file(file, allowed_extensions=None, max_size=None):
    """
    Validate uploaded file for security

    Args:
        file: FileStorage object from request.files
        allowed_extensions: Set of allowed extensions (None = all allowed types)
        max_size: Maximum file size in bytes (None = use tier-based limit)

    Returns:
        tuple: (is_valid: bool, error_message: str or None)
    """
    if not file or file.filename == "":
        return False, "No file selected"

    # Validate extension
    file_ext = Path(file.filename).suffix.lower()
    allowed = allowed_extensions or ALL_ALLOWED_EXTENSIONS

    if file_ext not in allowed:
        allowed_list = ", ".join(sorted(allowed))
        return False, f"File type '{file_ext}' not allowed. Allowed types: {allowed_list}"

    # Validate MIME type if provided
    if file.content_type and file.content_type not in ALLOWED_MIME_TYPES:
        return False, f"Invalid file MIME type: {file.content_type}"

    # Determine max size based on user tier if not specified
    if max_size is None:
        if current_user.is_authenticated:
            tier = current_user.tier
            if tier == TierLevel.PREMIUM:
                max_size = MAX_FILE_SIZE_PREMIUM
            elif tier == TierLevel.PROFESSIONAL:
                max_size = MAX_FILE_SIZE_PRO
            else:
                max_size = MAX_FILE_SIZE_FREE
        else:
            max_size = MAX_FILE_SIZE_FREE

    # Validate file size
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)  # Reset file pointer to start

    if file_size > max_size:
        max_size_mb = max_size / (1024 * 1024)
        actual_size_mb = file_size / (1024 * 1024)
        return (
            False,
            f"File too large ({actual_size_mb:.1f}MB). Maximum allowed: {max_size_mb:.0f}MB",
        )

    if file_size == 0:
        return False, "File is empty (0 bytes)"

    return True, None


# ========================================
# WEB ROUTES
# ========================================


@app.route("/")
def index():
    """Public landing page - Founding Member conversion focus"""
    try:
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        # New conversion-optimized landing page
        return render_template("landing-public.html")
    except Exception as e:
        # Fallback to a simple response if template fails
        app.logger.error(f"Index route error: {e}")
        return (
            f"""
        <html>
        <head><title>BarberX Legal Technologies</title></head>
        <body style="font-family: Arial; max-width: 800px; margin: 50px auto; padding: 20px;">
            <h1>BarberX Legal Technologies</h1>
            <p>Professional BWC Forensic Analysis Platform</p>
            <p><a href="/login">Login</a> | <a href="/auth/signup">Sign Up</a> | <a href="/health">Health Check</a></p>
            <p style="color: red;">Error: {e}</p>
        </body>
        </html>
        """,
            500,
        )


@app.route("/health")
def health():
    """Simple health check for Render"""
    return jsonify({"status": "ok", "timestamp": datetime.utcnow().isoformat()})


@app.route("/preview")
def preview_demo():
    """Free preview/demo page - no authentication required"""
    return render_template("preview-demo.html")


@app.route("/chat")
@login_required
def chat_interface():
    """Enhanced chat interface with memory and citations"""
    return render_template("chat/interface.html", user=current_user)


@app.route("/workspace")
@login_required
def unified_workspace():
    """Unified BarberX Workspace - All-in-one professional interface"""
    return render_template("unified-workspace.html", user=current_user)


# ========================================
# UNIFIED WORKSPACE API ENDPOINTS
# ========================================


@app.route("/api/workspace/models", methods=["GET"])
@login_required
def get_available_models():
    """Get list of available AI models"""
    models = [
        # OpenAI GPT-4o Series (Latest Flagship)
        {
            "id": "gpt-4o",
            "name": "GPT-4o",
            "provider": "openai",
            "description": "Fastest flagship model, multimodal capabilities",
            "context_window": 128000,
            "tier_required": "PRO",
        },
        {
            "id": "gpt-4o-mini",
            "name": "GPT-4o Mini",
            "provider": "openai",
            "description": "Affordable and intelligent small model",
            "context_window": 128000,
            "tier_required": "FREE",
        },
        # OpenAI GPT-4 Series
        {
            "id": "gpt-4-turbo",
            "name": "GPT-4 Turbo",
            "provider": "openai",
            "description": "Most capable GPT-4 model, best for complex analysis",
            "context_window": 128000,
            "tier_required": "PRO",
        },
        {
            "id": "gpt-4",
            "name": "GPT-4",
            "provider": "openai",
            "description": "Advanced reasoning and analysis",
            "context_window": 8192,
            "tier_required": "PRO",
        },
        {
            "id": "gpt-4-32k",
            "name": "GPT-4 32K",
            "provider": "openai",
            "description": "Extended context for large documents",
            "context_window": 32768,
            "tier_required": "PRO",
        },
        # OpenAI GPT-3.5 Series
        {
            "id": "gpt-3.5-turbo",
            "name": "GPT-3.5 Turbo",
            "provider": "openai",
            "description": "Fast and efficient for most tasks",
            "context_window": 16385,
            "tier_required": "FREE",
        },
        {
            "id": "gpt-3.5-turbo-16k",
            "name": "GPT-3.5 Turbo 16K",
            "provider": "openai",
            "description": "Extended context GPT-3.5",
            "context_window": 16385,
            "tier_required": "FREE",
        },
        # OpenAI o1 Series (Advanced Reasoning)
        {
            "id": "o1-preview",
            "name": "o1-preview",
            "provider": "openai",
            "description": "Advanced reasoning model for complex problems",
            "context_window": 128000,
            "tier_required": "ENTERPRISE",
        },
        {
            "id": "o1-mini",
            "name": "o1-mini",
            "provider": "openai",
            "description": "Faster reasoning for STEM tasks",
            "context_window": 128000,
            "tier_required": "PRO",
        },
        {
            "id": "o3-mini",
            "name": "o3-mini",
            "provider": "openai",
            "description": "Latest efficient reasoning model",
            "context_window": 128000,
            "tier_required": "PRO",
        },
        # Anthropic Claude 3.5 Series
        {
            "id": "claude-3-5-sonnet",
            "name": "Claude 3.5 Sonnet",
            "provider": "anthropic",
            "description": "Most intelligent Claude model, excellent for legal analysis",
            "context_window": 200000,
            "tier_required": "PRO",
        },
        # Anthropic Claude 3 Series
        {
            "id": "claude-3-opus",
            "name": "Claude 3 Opus",
            "provider": "anthropic",
            "description": "Powerful model for complex legal reasoning",
            "context_window": 200000,
            "tier_required": "ENTERPRISE",
        },
        {
            "id": "claude-3-sonnet",
            "name": "Claude 3 Sonnet",
            "provider": "anthropic",
            "description": "Balanced performance and speed",
            "context_window": 200000,
            "tier_required": "PRO",
        },
        {
            "id": "claude-3-haiku",
            "name": "Claude 3 Haiku",
            "provider": "anthropic",
            "description": "Fastest Claude model for quick analysis",
            "context_window": 200000,
            "tier_required": "FREE",
        },
        # GitHub Copilot Models
        {
            "id": "copilot-gpt-4",
            "name": "GitHub Copilot GPT-4",
            "provider": "github-copilot",
            "description": "Copilot-optimized GPT-4 for code and legal analysis",
            "context_window": 8192,
            "tier_required": "PRO",
        },
        {
            "id": "copilot-gpt-3.5-turbo",
            "name": "GitHub Copilot GPT-3.5",
            "provider": "github-copilot",
            "description": "Fast Copilot model for quick insights",
            "context_window": 16385,
            "tier_required": "FREE",
        },
        # Google Gemini Series
        {
            "id": "gemini-pro",
            "name": "Gemini Pro",
            "provider": "google",
            "description": "Google's advanced AI for complex reasoning",
            "context_window": 32000,
            "tier_required": "PRO",
        },
        {
            "id": "gemini-pro-vision",
            "name": "Gemini Pro Vision",
            "provider": "google",
            "description": "Multimodal model for image and text analysis",
            "context_window": 32000,
            "tier_required": "PRO",
        },
        # Microsoft Azure OpenAI
        {
            "id": "azure-gpt-4",
            "name": "Azure GPT-4",
            "provider": "azure-openai",
            "description": "Enterprise-grade GPT-4 via Azure",
            "context_window": 8192,
            "tier_required": "ENTERPRISE",
        },
        {
            "id": "azure-gpt-4-turbo",
            "name": "Azure GPT-4 Turbo",
            "provider": "azure-openai",
            "description": "Enterprise GPT-4 Turbo with enhanced security",
            "context_window": 128000,
            "tier_required": "ENTERPRISE",
        },
    ]

    # Filter by user tier
    user_tier = current_user.tier.value if current_user.tier else "FREE"
    tier_hierarchy = {"FREE": 0, "PRO": 1, "ENTERPRISE": 2}
    user_tier_level = tier_hierarchy.get(user_tier, 0)

    accessible_models = [
        m for m in models if tier_hierarchy.get(m["tier_required"], 2) <= user_tier_level
    ]

    return jsonify({"models": accessible_models})


@app.route("/api/workspace/tools", methods=["GET"])
@login_required
def get_available_tools():
    """Get list of available workspace tools"""
    tools = [
        {
            "id": "chat",
            "name": "AI Chat",
            "icon": "comments",
            "description": "Intelligent legal assistant",
            "tier_required": "FREE",
        },
        {
            "id": "analyzer",
            "name": "Evidence Analyzer",
            "icon": "microscope",
            "description": "BWC and document analysis",
            "tier_required": "PRO",
        },
        {
            "id": "timeline",
            "name": "Timeline Builder",
            "icon": "clock",
            "description": "Event chronology tool",
            "tier_required": "PRO",
        },
        {
            "id": "discrepancy",
            "name": "Discrepancy Detector",
            "icon": "search",
            "description": "Find inconsistencies",
            "tier_required": "PRO",
        },
        {
            "id": "transcript",
            "name": "Transcript Analyzer",
            "icon": "file-alt",
            "description": "Transcription and analysis",
            "tier_required": "PRO",
        },
        {
            "id": "workflow",
            "name": "Custom Workflows",
            "icon": "project-diagram",
            "description": "Automated legal workflows",
            "tier_required": "ENTERPRISE",
        },
    ]

    return jsonify({"tools": tools})


@app.route("/api/workspace/analyze", methods=["POST"])
@login_required
@require_tier(TierLevel.PROFESSIONAL)
def analyze_evidence():
    """Unified evidence analysis endpoint"""
    import time

    from usage_meter import SmartMeter, UsageQuota

    # Check quota BEFORE processing
    quota = UsageQuota.query.filter_by(user_id=current_user.id).first()
    if not quota:
        quota = SmartMeter.initialize_user_quota(current_user.id)

    has_quota, error_msg = quota.check_quota("analyses")
    if not has_quota:
        SmartMeter.track_event(
            event_type="analysis_blocked",
            event_category="quota",
            status="denied",
            error_message=error_msg,
        )
        return (
            jsonify({"error": "Quota exceeded", "message": error_msg, "upgrade_url": "/pricing"}),
            429,
        )

    data = request.get_json()
    analysis_type = data.get("type", "general")
    content = data.get("content", "")
    model = data.get("model", "gpt-3.5-turbo")

    if not content:
        return jsonify({"error": "No content provided"}), 400

    start_time = time.time()

    try:
        # Use existing ChatGPT service
        from chatgpt_service import ChatGPTService

        # Get user's API key or use system key
        user_api_key = session.get("openai_api_key") or os.getenv("OPENAI_API_KEY")

        if not user_api_key:
            return (
                jsonify(
                    {
                        "error": "No API key configured",
                        "message": "Please add your OpenAI API key in settings",
                    }
                ),
                400,
            )

        chatgpt = ChatGPTService(api_key=user_api_key)

        # Build analysis prompt based on type
        prompts = {
            "general": "Analyze this legal evidence and provide key insights:",
            "timeline": "Extract and organize all events in chronological order:",
            "discrepancy": "Identify any inconsistencies, contradictions, or discrepancies:",
            "summary": "Provide a concise summary of this evidence:",
            "transcript": "Analyze this transcript for key statements and themes:",
        }

        system_prompt = prompts.get(analysis_type, prompts["general"])

        messages = [
            {
                "role": "system",
                "content": "You are an expert legal analyst specializing in evidence review and BWC analysis.",
            },
            {"role": "user", "content": f"{system_prompt}\n\n{content}"},
        ]

        response = chatgpt.create_chat_completion(
            messages=messages, model=model, max_tokens=2000, temperature=0.3
        )

        duration = time.time() - start_time
        tokens_used = response.get("tokens_used", 0)

        # Estimate cost (approximate rates)
        cost_per_1k = 0.002 if "gpt-4" in model else 0.0005
        estimated_cost = (tokens_used / 1000) * cost_per_1k

        if response.get("success"):
            # Track successful analysis
            SmartMeter.track_event(
                event_type="analysis",
                event_category="compute",
                resource_name=model,
                tokens_input=len(content.split()),
                tokens_output=tokens_used,
                duration_seconds=duration,
                cost_usd=estimated_cost,
                status="success",
            )

            # Increment quota
            quota.increment_quota("analyses")
            quota.increment_quota("ai_tokens", tokens_used)
            quota.total_cost_usd += estimated_cost
            db.session.commit()

            return jsonify(
                {
                    "analysis": response["content"],
                    "model": response.get("model"),
                    "tokens_used": response.get("tokens_used"),
                }
            )
        else:
            # Track failed analysis
            SmartMeter.track_event(
                event_type="analysis",
                event_category="compute",
                resource_name=model,
                duration_seconds=duration,
                status="error",
                error_message=response.get("error"),
            )
            return jsonify({"error": response.get("error")}), 500

    except Exception as e:
        app.logger.error(f"Analysis error: {e}")
        # Track error
        SmartMeter.track_event(
            event_type="analysis",
            event_category="compute",
            resource_name=model,
            status="error",
            error_message=str(e),
        )
        return jsonify({"error": str(e)}), 500


@app.route("/api/workspace/workflows", methods=["GET"])
@login_required
def get_workflows():
    """Get available custom workflows"""
    workflows = [
        {
            "id": "brady-analysis",
            "name": "Brady Violation Analysis",
            "description": "Systematic analysis for exculpatory evidence",
            "steps": [
                "Extract all evidence items",
                "Identify potentially exculpatory material",
                "Check disclosure timeline",
                "Generate Brady checklist",
            ],
            "tier_required": "PRO",
        },
        {
            "id": "bwc-forensic",
            "name": "BWC Forensic Workflow",
            "description": "Complete body camera analysis pipeline",
            "steps": [
                "Metadata extraction",
                "Video transcript generation",
                "Scene analysis",
                "Officer statement comparison",
                "Timeline reconstruction",
            ],
            "tier_required": "PRO",
        },
        {
            "id": "discovery-review",
            "name": "Discovery Document Review",
            "description": "Automated document classification and analysis",
            "steps": [
                "Document classification",
                "Entity extraction",
                "Key fact identification",
                "Privilege review",
                "Summary generation",
            ],
            "tier_required": "ENTERPRISE",
        },
    ]

    return jsonify({"workflows": workflows})


@app.route("/api/workspace/execute-workflow", methods=["POST"])
@login_required
@require_tier(TierLevel.PROFESSIONAL)
def execute_workflow():
    """Execute a custom workflow"""
    data = request.get_json()
    workflow_id = data.get("workflow_id")
    input_data = data.get("input")

    if not workflow_id or not input_data:
        return jsonify({"error": "Missing workflow_id or input"}), 400

    # Workflow execution logic here
    # This would integrate with various analysis services

    return jsonify(
        {
            "status": "processing",
            "workflow_id": workflow_id,
            "message": "Workflow execution started",
            "estimated_time": "2-5 minutes",
        }
    )


# ========================================
# SMART METER USAGE TRACKING API
# ========================================


@app.route("/api/usage/stats", methods=["GET"])
@login_required
def get_usage_stats():
    """Get comprehensive usage statistics for current user"""
    from usage_meter import SmartMeter

    days = request.args.get("days", 30, type=int)
    stats = SmartMeter.get_user_stats(current_user.id, days=days)

    return jsonify(stats)


@app.route("/api/usage/quota", methods=["GET"])
@login_required
def get_usage_quota():
    """Get real-time quota status for current user"""
    from usage_meter import SmartMeter, UsageQuota

    quota = UsageQuota.query.filter_by(user_id=current_user.id).first()
    if not quota:
        quota = SmartMeter.initialize_user_quota(current_user.id)

    quota.reset_if_new_period()

    return jsonify(
        {
            "period": {
                "start": quota.period_start.isoformat(),
                "end": quota.period_end.isoformat(),
                "days_remaining": (quota.period_end - datetime.utcnow()).days,
            },
            "quotas": {
                "ai_tokens": {
                    "used": quota.ai_tokens_used,
                    "limit": quota.ai_tokens_limit,
                    "percent": quota.get_usage_percent("ai_tokens"),
                    "remaining": (
                        max(0, quota.ai_tokens_limit - quota.ai_tokens_used)
                        if quota.ai_tokens_limit != -1
                        else -1
                    ),
                },
                "ai_requests": {
                    "used": quota.ai_requests_count,
                    "limit": quota.ai_requests_limit,
                    "percent": quota.get_usage_percent("ai_requests"),
                    "remaining": (
                        max(0, quota.ai_requests_limit - quota.ai_requests_count)
                        if quota.ai_requests_limit != -1
                        else -1
                    ),
                },
                "storage": {
                    "used": quota.storage_bytes_used,
                    "limit": quota.storage_bytes_limit,
                    "percent": quota.get_usage_percent("storage"),
                    "remaining": (
                        max(0, quota.storage_bytes_limit - quota.storage_bytes_used)
                        if quota.storage_bytes_limit != -1
                        else -1
                    ),
                    "used_mb": round(quota.storage_bytes_used / 1048576, 2),
                    "limit_mb": (
                        round(quota.storage_bytes_limit / 1048576, 2)
                        if quota.storage_bytes_limit != -1
                        else -1
                    ),
                },
                "files": {
                    "used": quota.files_uploaded_count,
                    "limit": quota.files_uploaded_limit,
                    "percent": quota.get_usage_percent("files"),
                    "remaining": (
                        max(0, quota.files_uploaded_limit - quota.files_uploaded_count)
                        if quota.files_uploaded_limit != -1
                        else -1
                    ),
                },
                "analyses": {
                    "used": quota.analyses_count,
                    "limit": quota.analyses_limit,
                    "percent": quota.get_usage_percent("analyses"),
                    "remaining": (
                        max(0, quota.analyses_limit - quota.analyses_count)
                        if quota.analyses_limit != -1
                        else -1
                    ),
                },
                "workflows": {
                    "used": quota.workflows_executed_count,
                    "limit": quota.workflows_executed_limit,
                    "percent": quota.get_usage_percent("workflows"),
                    "remaining": (
                        max(0, quota.workflows_executed_limit - quota.workflows_executed_count)
                        if quota.workflows_executed_limit != -1
                        else -1
                    ),
                },
                "api_calls": {
                    "used": quota.api_calls_count,
                    "limit": quota.api_calls_limit,
                    "percent": quota.get_usage_percent("api_calls"),
                    "remaining": (
                        max(0, quota.api_calls_limit - quota.api_calls_count)
                        if quota.api_calls_limit != -1
                        else -1
                    ),
                },
                "cost": {
                    "used_usd": float(quota.total_cost_usd),
                    "limit_usd": float(quota.cost_limit_usd),
                    "percent": (
                        min(
                            100.0, (float(quota.total_cost_usd) / float(quota.cost_limit_usd)) * 100
                        )
                        if quota.cost_limit_usd > 0
                        else 0
                    ),
                    "remaining_usd": (
                        max(0, float(quota.cost_limit_usd) - float(quota.total_cost_usd))
                        if quota.cost_limit_usd != -1
                        else -1
                    ),
                },
            },
            "alerts": {
                "alert_80_percent": quota.alert_80_percent_sent,
                "alert_95_percent": quota.alert_95_percent_sent,
                "alert_100_percent": quota.alert_100_percent_sent,
            },
        }
    )


@app.route("/api/usage/events", methods=["GET"])
@login_required
def get_usage_events():
    """Get recent usage events for current user"""
    from usage_meter import SmartMeterEvent

    limit = request.args.get("limit", 100, type=int)
    event_type = request.args.get("type")
    days = request.args.get("days", 7, type=int)

    since = datetime.utcnow() - timedelta(days=days)

    query = SmartMeterEvent.query.filter(
        SmartMeterEvent.user_id == current_user.id, SmartMeterEvent.timestamp >= since
    )

    if event_type:
        query = query.filter(SmartMeterEvent.event_type == event_type)

    events = query.order_by(SmartMeterEvent.timestamp.desc()).limit(limit).all()

    return jsonify(
        {
            "events": [
                {
                    "id": e.id,
                    "event_type": e.event_type,
                    "event_category": e.event_category,
                    "resource_name": e.resource_name,
                    "quantity": e.quantity,
                    "tokens_input": e.tokens_input,
                    "tokens_output": e.tokens_output,
                    "duration_seconds": e.duration_seconds,
                    "file_size_bytes": e.file_size_bytes,
                    "cost_usd": float(e.cost_usd) if e.cost_usd else 0,
                    "status": e.status,
                    "endpoint": e.endpoint,
                    "timestamp": e.timestamp.isoformat(),
                }
                for e in events
            ],
            "total": len(events),
        }
    )


@app.route("/api/usage/summary", methods=["GET"])
@login_required
def get_usage_summary():
    """Get usage summary with charts data"""
    from sqlalchemy import func

    from usage_meter import SmartMeter, SmartMeterEvent

    days = request.args.get("days", 30, type=int)
    since = datetime.utcnow() - timedelta(days=days)

    # Get daily breakdown
    daily_stats = (
        db.session.query(
            func.date(SmartMeterEvent.timestamp).label("date"),
            func.count(SmartMeterEvent.id).label("count"),
            func.sum(SmartMeterEvent.tokens_input + SmartMeterEvent.tokens_output).label("tokens"),
            func.sum(SmartMeterEvent.cost_usd).label("cost"),
        )
        .filter(SmartMeterEvent.user_id == current_user.id, SmartMeterEvent.timestamp >= since)
        .group_by(func.date(SmartMeterEvent.timestamp))
        .order_by("date")
        .all()
    )

    # Get by event type
    by_type = (
        db.session.query(
            SmartMeterEvent.event_type,
            func.count(SmartMeterEvent.id).label("count"),
            func.sum(SmartMeterEvent.cost_usd).label("cost"),
        )
        .filter(SmartMeterEvent.user_id == current_user.id, SmartMeterEvent.timestamp >= since)
        .group_by(SmartMeterEvent.event_type)
        .all()
    )

    # Get comprehensive stats
    stats = SmartMeter.get_user_stats(current_user.id, days=days)

    return jsonify(
        {
            "daily": [
                {
                    "date": str(row.date),
                    "events": row.count,
                    "tokens": int(row.tokens or 0),
                    "cost_usd": float(row.cost or 0),
                }
                for row in daily_stats
            ],
            "by_type": [
                {
                    "type": row.event_type,
                    "count": row.count,
                    "cost_usd": float(row.cost or 0),
                }
                for row in by_type
            ],
            "stats": stats,
        }
    )


@app.route("/api/usage/track", methods=["POST"])
@login_required
def track_usage_event():
    """Manually track a usage event (for client-side tracking)"""
    from usage_meter import SmartMeter, UsageQuota

    data = request.get_json()

    # Check rate limit
    quota = UsageQuota.query.filter_by(user_id=current_user.id).first()
    if not quota:
        quota = SmartMeter.initialize_user_quota(current_user.id)

    if not quota.check_rate_limit():
        return (
            jsonify(
                {
                    "error": "Rate limit exceeded",
                    "message": "Too many requests. Please slow down.",
                    "retry_after": 60,
                }
            ),
            429,
        )

    # Update rate limit counter
    quota.requests_this_minute += 1
    quota.last_request_timestamp = datetime.utcnow()
    db.session.commit()

    event = SmartMeter.track_event(
        event_type=data.get("event_type", "unknown"),
        event_category=data.get("event_category", "feature"),
        user_id=current_user.id,
        resource_name=data.get("resource_name"),
        quantity=data.get("quantity", 1.0),
        duration_seconds=data.get("duration_seconds", 0.0),
        status=data.get("status", "success"),
    )

    return jsonify(
        {
            "success": True,
            "event_id": event.id if event else None,
        }
    )


# ========================================
# BACKEND OPTIMIZATION API ENDPOINTS
# ========================================


@app.route("/health-detailed")
def health_check_detailed():
    """Detailed system health check endpoint"""
    if not BACKEND_OPTIMIZATION_AVAILABLE:
        return jsonify(
            {
                "status": "healthy",
                "timestamp": datetime.utcnow().isoformat(),
                "backend_optimization": "not available",
            }
        )

    # Check database
    db_healthy = True
    try:
        db.session.execute("SELECT 1")
    except Exception:
        db_healthy = False

    # Check services
    services = service_registry.list_services()
    services_healthy = all(s["status"] == "active" for s in services)

    # Get performance stats
    from backend_integration import performance_monitor

    stats = performance_monitor.get_stats()

    health_status = {
        "status": "healthy" if (db_healthy and services_healthy) else "degraded",
        "timestamp": datetime.utcnow().isoformat(),
        "components": {
            "database": "up" if db_healthy else "down",
            "services": "up" if services_healthy else "degraded",
        },
        "metrics": {
            "total_requests": sum(s.get("call_count", 0) for s in stats.values()),
            "registered_services": len(services),
        },
    }

    status_code = 200 if health_status["status"] == "healthy" else 503

    return jsonify(health_status), status_code


@app.route("/api/rate-limit/status")
def api_rate_limit_status():
    """Get current rate limit status for user"""
    if not BACKEND_OPTIMIZATION_AVAILABLE:
        return jsonify(error_response("Backend optimization not available")), 503

    from api_middleware import rate_limiter

    if current_user.is_authenticated:
        identifier = str(current_user.id)
        tier = current_user.tier
    else:
        identifier = request.remote_addr
        tier = "free"

    remaining = rate_limiter.get_remaining(identifier, tier)
    limit = rate_limiter.tier_limits[tier]

    return jsonify(
        success_response(
            "Rate limit status",
            {
                "tier": tier,
                "limit_per_minute": limit,
                "remaining": remaining,
                "reset_in_seconds": 60,
            },
        )
    )


@app.route("/bwc-dashboard")
@login_required
def bwc_dashboard():
    """BWC Analysis Dashboard - Frontend Interface"""
    return render_template("bwc-dashboard.html", user=current_user)


@app.route("/test-separation")
def test_separation():
    """Frontend/Backend separation test suite"""
    return send_file("test_separation.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    """User registration - redirect to enhanced signup"""
    if ENHANCED_AUTH_AVAILABLE:
        return redirect(url_for("auth.signup"))

    # Fallback to old registration logic
    if request.method == "GET":
        return send_file("templates/register.html")

    from utils.logging_config import get_logger
    from utils.responses import (error_response, success_response,
                                 validation_error)
    from utils.security import InputValidator

    logger = get_logger("auth")

    try:
        data = request.get_json()

        # Validate required fields
        validation_errors = {}

        email = data.get("email", "").strip()
        password = data.get("password", "")
        full_name = data.get("full_name", "").strip()

        # Email validation
        if not email:
            validation_errors["email"] = ["Email is required"]
        else:
            is_valid, error_msg = InputValidator.validate_email(email)
            if not is_valid:
                validation_errors["email"] = [error_msg]

        # Password validation
        if not password:
            validation_errors["password"] = ["Password is required"]
        else:
            is_valid, error_msg = InputValidator.validate_password(password)
            if not is_valid:
                validation_errors["password"] = [error_msg]

        # Name validation
        if not full_name:
            validation_errors["full_name"] = ["Full name is required"]
        elif len(full_name) > 100:
            validation_errors["full_name"] = ["Name too long (max 100 characters)"]

        # Return validation errors if any
        if validation_errors:
            return validation_error(validation_errors)

        # Check if user exists
        if User.query.filter_by(email=email).first():
            return error_response(
                "This email is already registered", error_code="ALREADY_EXISTS", status_code=400
            )

        # Create user
        user = User(
            email=email,
            full_name=InputValidator.sanitize_text(full_name, 100),
            organization=InputValidator.sanitize_text(data.get("organization", ""), 100),
            subscription_tier="free",
        )
        user.set_password(password)

        db.session.add(user)
        db.session.commit()

        # Log audit
        AuditLog.log("user_registered", "user", str(user.id))

        # Auto-login
        login_user(user)

        logger.info(f"New user registered: {user.email}")

        # Redirect new users to onboarding welcome screen
        return success_response(
            data={"user": user.to_dict(), "redirect": "/welcome"},
            message="Registration successful",
            status_code=201,
        )

    except Exception as e:
        logger.error(f"Registration failed: {type(e).__name__}: {e}", exc_info=True)
        from utils.security import ErrorSanitizer

        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            "Registration failed. Please try again.",
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


# [SECURITY FIX] Old /login route removed - enhanced /auth/login has proper redirect protection
# Enhanced auth route at /auth/login includes is_safe_url() validation to prevent open redirects


@app.route("/logout")
@login_required
def logout():
    """User logout"""
    AuditLog.log("user_logout", "user", str(current_user.id))
    logout_user()
    return redirect(url_for("index"))


@app.route("/batch-pdf-upload.html")
@login_required
def batch_pdf_upload():
    """Batch PDF upload page"""
    return render_template("batch-pdf-upload.html", user=current_user)


@app.route("/welcome")
@login_required
def welcome():
    """First-time user onboarding welcome screen"""
    return render_template("onboarding/welcome.html", user=current_user)


@app.route("/skip-onboarding", methods=["POST"])
@login_required
def skip_onboarding():
    """Mark onboarding as complete and redirect to dashboard"""
    try:
        # Mark onboarding complete in session (database update can be added later)
        session["onboarding_complete"] = True

        from utils.logging_config import get_logger

        logger = get_logger("app")
        logger.info(f"User {current_user.email} skipped onboarding")

        return redirect(url_for("dashboard"))
    except Exception as e:
        app.logger.error(f"Skip onboarding error: {e}")
        return redirect(url_for("dashboard"))


@app.route("/dashboard")
@login_required
def dashboard():
    """Enhanced user dashboard with usage tracking and tier-specific features"""
    try:
        from models_auth import UsageTracking

        usage = UsageTracking.get_or_create_current(current_user.id)
        limits = current_user.get_tier_limits()

        return render_template("auth/dashboard.html", user=current_user, usage=usage, limits=limits)
    except Exception as e:
        app.logger.error(f"Dashboard error: {e}", exc_info=True)
        # Create minimal usage/limits on error
        usage = type(
            "obj",
            (object,),
            {
                "bwc_videos_processed": 0,
                "document_pages_processed": 0,
                "transcription_minutes_used": 0,
                "storage_used_mb": 0,
            },
        )()
        limits = {
            "bwc_videos_per_month": 1,
            "document_pages_per_month": 5,
            "transcription_minutes_per_month": 10,
            "storage_gb": 1,
        }
        return render_template("auth/dashboard.html", user=current_user, usage=usage, limits=limits)


@app.route("/account-settings")
@login_required
def account_settings_alias():
    """Alias for account settings page"""
    return redirect(url_for("account_settings"))


@app.route("/admin")
@login_required
def admin_panel():
    """Enhanced admin panel - requires admin role with full analytics"""
    if not hasattr(current_user, "is_admin") or not current_user.is_admin:
        flash("Admin access required", "danger")
        return redirect(url_for("dashboard"))

    return send_file("templates/admin/admin-dashboard-enhanced.html")


@app.route("/admin/founding-members")
@login_required
def admin_founding_members():
    """Admin view of founding member signups"""
    if not hasattr(current_user, "is_admin") or not current_user.is_admin:
        flash("Admin access required", "danger")
        return redirect(url_for("dashboard"))

    import csv

    signups_file = Path("founding_member_signups.csv")
    signups = []

    if signups_file.exists():
        with open(signups_file, "r", newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            signups = list(reader)

    spots_remaining = max(0, 100 - len(signups))

    return render_template(
        "admin/founding-members.html",
        signups=signups,
        total_signups=len(signups),
        spots_remaining=spots_remaining,
    )


@app.route("/admin/founding-members/export")
@login_required
def admin_export_founding_members():
    """Export founding member signups as CSV"""
    if not hasattr(current_user, "is_admin") or not current_user.is_admin:
        return jsonify({"error": "Admin access required"}), 403

    signups_file = Path("founding_member_signups.csv")

    if not signups_file.exists():
        return jsonify({"error": "No signups found"}), 404

    return send_file(
        signups_file,
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"founding_members_{datetime.utcnow().strftime('%Y%m%d')}.csv",
    )


@app.route("/account")
@login_required
def account_settings():
    """User account settings page"""
    if ENHANCED_AUTH_AVAILABLE:
        from models_auth import UsageTracking

        # Get usage for current month
        usage = UsageTracking.query.filter_by(
            user_id=current_user.id, month=datetime.utcnow().month, year=datetime.utcnow().year
        ).first()

        # Get tier limits
        limits = current_user.get_tier_limits()

        return render_template(
            "auth/account-settings.html",
            usage=usage or UsageTracking(user_id=current_user.id),
            limits=limits,
        )
    else:
        return send_file("templates/auth/account-settings.html")


@app.route("/api/user/profile", methods=["PUT"])
@login_required
def update_user_profile():
    """Update user profile information"""
    if not ENHANCED_AUTH_AVAILABLE:
        return jsonify({"error": "Feature not available"}), 503

    from models_auth import db

    try:
        data = request.get_json()

        if data.get("full_name"):
            current_user.full_name = data["full_name"]
        if data.get("organization"):
            current_user.organization = data["organization"]

        db.session.commit()

        return jsonify({"message": "Profile updated successfully"})
    except Exception as e:
        app.logger.error(f"Profile update error: {e}")
        return jsonify({"error": "Failed to update profile"}), 400


@app.route("/api/user/change-password", methods=["POST"])
@login_required
def change_password():
    """Change user password"""
    if not ENHANCED_AUTH_AVAILABLE:
        return jsonify({"error": "Feature not available"}), 503

    from models_auth import db
    from utils.logging_config import get_logger
    from utils.responses import (error_response, success_response,
                                 validation_error)
    from utils.security import ErrorSanitizer, InputValidator

    logger = get_logger("auth")

    try:
        data = request.get_json()
        current_password = data.get("current_password", "")
        new_password = data.get("new_password", "")

        # Validate new password
        is_valid, error_msg = InputValidator.validate_password(new_password)
        if not is_valid:
            return validation_error({"new_password": [error_msg]})

        # Check current password
        if not current_user.check_password(current_password):
            return error_response(
                "Current password is incorrect", error_code="INVALID_CREDENTIALS", status_code=401
            )

        # Update password
        current_user.set_password(new_password)
        db.session.commit()

        logger.info(f"Password changed for user: {current_user.email}")

        return success_response(message="Password changed successfully")

    except Exception as e:
        logger.error(f"Password change failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            "Failed to change password",
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/user/export-data", methods=["POST"])
@login_required
def export_user_data():
    """Request user data export (GDPR compliance)"""
    # In production, this would trigger an async job to compile user data
    app.logger.info(f"Data export requested by user {current_user.id}")
    return jsonify({"message": "Export will be emailed when ready"})


@app.route("/api/user/delete-account", methods=["DELETE"])
@login_required
def delete_user_account():
    """Delete user account (GDPR compliance)"""
    if not ENHANCED_AUTH_AVAILABLE:
        return jsonify({"error": "Feature not available"}), 503

    from flask_login import logout_user

    from models_auth import db

    try:
        user_id = current_user.id
        db.session.delete(current_user)
        db.session.commit()
        logout_user()

        app.logger.info(f"User account {user_id} deleted")
        return jsonify({"message": "Account deleted"})
    except Exception as e:
        app.logger.error(f"Account deletion error: {e}")
        return jsonify({"error": "Failed to delete account"}), 400


@app.route("/admin")
@login_required
def admin_panel_old():
    """Enhanced admin panel - requires admin role with full analytics"""
    if not hasattr(current_user, "role") or current_user.role != "admin":
        flash("Admin access required", "danger")
        return redirect(url_for("dashboard"))

    if ENHANCED_AUTH_AVAILABLE:
        try:
            from sqlalchemy import func

            from models_auth import UsageTracking, User

            # Optimize: Get aggregated stats without loading all users
            total_users = User.query.count()
            total_analyses = (
                db.session.query(func.sum(UsageTracking.bwc_videos_processed)).scalar() or 0
            )
            total_storage = db.session.query(func.sum(UsageTracking.storage_used_mb)).scalar() or 0
            storage_gb = round(total_storage / 1024, 2) if total_storage else 0

            # Calculate MRR efficiently using SQL aggregation
            from models_auth import TierLevel

            revenue = 0
            for tier in TierLevel:
                tier_count = User.query.filter_by(tier=tier).count()
                revenue += tier_count * tier.value

            # Only get users for display (with limit)
            users = User.query.order_by(User.created_at.desc()).limit(100).all()

            return render_template(
                "admin/dashboard.html",
                users=users,
                total_users=total_users,
                total_analyses=total_analyses,
                storage_gb=storage_gb,
                revenue=revenue,
            )
        except Exception as e:
            app.logger.error(f"Admin panel error: {e}")
            return send_file("admin.html")
    else:
        return send_file("admin.html")


@app.route("/analyzer")
@login_required
def analyzer():
    """BWC analyzer interface"""
    return send_file("bwc-analyzer.html")


@app.route("/analysis/<analysis_id>")
@login_required
def view_analysis_results(analysis_id):
    """View enhanced analysis results page"""
    return send_file("templates/enhanced-analysis.html")


@app.route("/api/analysis/<analysis_id>/enhanced", methods=["GET"])
@login_required
def get_enhanced_analysis(analysis_id):
    """Get enhanced analysis data with all features"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    # Load the full report from file
    if analysis.report_json_path and os.path.exists(analysis.report_json_path):
        import json

        with open(analysis.report_json_path, "r") as f:
            report_data = json.load(f)
        return jsonify(report_data)

    return jsonify({"error": "Analysis report not available"}), 404


@app.route("/api/analysis/<analysis_id>/video", methods=["GET"])
@login_required
def get_analysis_video(analysis_id):
    """Stream video file for analysis"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    if not analysis.file_path or not os.path.exists(analysis.file_path):
        return jsonify({"error": "Video file not found"}), 404

    return send_file(analysis.file_path, mimetype="video/mp4")


@app.route("/api/analysis/<analysis_id>/report/<format>", methods=["GET"])
@login_required
def download_analysis_report(analysis_id, format):
    """Download analysis report in specified format"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    format = format.lower()

    if format == "json" and analysis.report_json_path:
        return send_file(
            analysis.report_json_path,
            as_attachment=True,
            download_name=f"{analysis.id}_report.json",
        )
    elif format == "txt" and analysis.report_txt_path:
        return send_file(
            analysis.report_txt_path, as_attachment=True, download_name=f"{analysis.id}_report.txt"
        )
    elif format == "md" and analysis.report_md_path and os.path.exists(analysis.report_md_path):
        return send_file(
            analysis.report_md_path, as_attachment=True, download_name=f"{analysis.id}_report.md"
        )

    return jsonify({"error": f"Report format '{format}' not available"}), 404


@app.route("/api/analyses", methods=["GET"])
@login_required
def get_user_analyses():
    """Get all analyses for current user"""
    analyses = (
        Analysis.query.filter_by(user_id=current_user.id).order_by(Analysis.created_at.desc()).all()
    )

    return jsonify(
        {
            "analyses": [
                {
                    "id": a.id,
                    "filename": a.filename,
                    "status": a.status,
                    "case_number": a.case_number,
                    "created_at": a.created_at.isoformat() if a.created_at else None,
                    "progress": a.progress,
                }
                for a in analyses
            ]
        }
    )


@app.route("/api/health", methods=["GET"])
def health_check():
    """System health check"""
    return jsonify({"status": "ok", "timestamp": datetime.utcnow().isoformat()})


# Evidence Processing Routes
@app.route("/evidence/intake")
@login_required
def evidence_intake_page():
    """Evidence intake form"""
    return send_file("templates/evidence-intake.html")


@app.route("/evidence/dashboard")
@login_required
def evidence_dashboard_page():
    """Evidence processing dashboard"""
    return send_file("templates/evidence-dashboard.html")


@app.route("/api/evidence/intake", methods=["POST"])
@login_required
def evidence_intake_submit():
    """Submit new evidence for processing"""
    from evidence_processing import evidence_workflow
    from utils.logging_config import get_logger
    from utils.responses import (error_response, success_response,
                                 validation_error)
    from utils.security import ErrorSanitizer, InputValidator

    logger = get_logger("api")

    try:
        # Validate required fields
        validation_errors = {}

        case_number = request.form.get("case_number", "").strip()
        if not case_number:
            validation_errors["case_number"] = ["Case number is required"]
        elif len(case_number) > 50:
            validation_errors["case_number"] = ["Case number too long (max 50 characters)"]

        evidence_type = request.form.get("evidence_type", "").strip()
        if not evidence_type:
            validation_errors["evidence_type"] = ["Evidence type is required"]

        # Return validation errors if any
        if validation_errors:
            return validation_error(validation_errors)

        # Get form data with sanitization
        data = {
            "case_number": InputValidator.sanitize_text(case_number, 50),
            "incident_date": request.form.get("incident_date", "").strip(),
            "incident_location": InputValidator.sanitize_text(
                request.form.get("incident_location", ""), 200
            ),
            "case_type": request.form.get("case_type", "").strip(),
            "jurisdiction": InputValidator.sanitize_text(request.form.get("jurisdiction", ""), 100),
            "lead_investigator": InputValidator.sanitize_text(
                request.form.get("lead_investigator", ""), 100
            ),
            "evidence_type": evidence_type,
            "description": InputValidator.sanitize_text(request.form.get("description", ""), 5000),
            "source": InputValidator.sanitize_text(request.form.get("source", ""), 200),
            "officer_name": InputValidator.sanitize_text(request.form.get("officer_name", ""), 100),
            "badge_number": InputValidator.sanitize_text(request.form.get("badge_number", ""), 50),
            "acquired_by": InputValidator.sanitize_text(request.form.get("acquired_by", ""), 100),
            "acquired_date": request.form.get("acquired_date", "").strip(),
            "acquisition_method": InputValidator.sanitize_text(
                request.form.get("acquisition_method", ""), 200
            ),
            "storage_location": InputValidator.sanitize_text(
                request.form.get("storage_location", ""), 200
            ),
            "priority": request.form.get("priority", "normal"),
            "assigned_to": InputValidator.sanitize_text(request.form.get("assigned_to", ""), 100),
            "special_instructions": InputValidator.sanitize_text(
                request.form.get("special_instructions", ""), 2000
            ),
            "submitted_by": current_user.email,
            "id": str(uuid.uuid4())[:12].upper(),
        }

        # Parse tags safely
        try:
            tags_str = request.form.get("tags", "[]")
            data["tags"] = json.loads(tags_str) if tags_str else []
        except json.JSONDecodeError:
            data["tags"] = []

        # Handle file upload with validation
        if "files" in request.files:
            files = request.files.getlist("files")
            for file in files:
                if file and file.filename:
                    # Validate file type
                    is_valid, error_msg = InputValidator.validate_file_type(file)
                    if not is_valid:
                        logger.warning(f"File upload rejected: {error_msg}")
                        return error_response(
                            error_msg, error_code="FILE_TYPE_NOT_ALLOWED", status_code=400
                        )

                    # Validate file size (get category from evidence type)
                    category = "video" if "video" in evidence_type.lower() else "document"
                    is_valid, error_msg = InputValidator.validate_file_size(file, category)
                    if not is_valid:
                        logger.warning(f"File upload rejected: {error_msg}")
                        return error_response(
                            error_msg, error_code="FILE_TOO_LARGE", status_code=400
                        )

                    # Sanitize filename and save securely
                    filename = secure_filename(file.filename)
                    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                    unique_filename = f"{data['id']}_{timestamp}_{filename}"

                    # Save file with secure path
                    upload_dir = Path(app.config.get("UPLOAD_FOLDER", "./uploads/evidence"))
                    upload_dir.mkdir(parents=True, exist_ok=True)

                    try:
                        filepath = InputValidator.sanitize_path(str(upload_dir), unique_filename)
                    except ValueError as e:
                        logger.error(f"Path traversal attempt detected: {e}")
                        return error_response(
                            "Invalid file path", error_code="VALIDATION_ERROR", status_code=400
                        )

                    file.save(filepath)
                    logger.info(f"File saved: {filepath}")

                    # Calculate hash for integrity
                    file_hash = hashlib.sha256()
                    with open(filepath, "rb") as f:
                        for chunk in iter(lambda: f.read(8192), b""):
                            file_hash.update(chunk)

                    data["filename"] = filename
                    data["file_path"] = str(filepath)
                    data["file_size"] = os.path.getsize(filepath)
                    data["file_hash"] = file_hash.hexdigest()
                    data["format"] = filename.split(".")[-1].lower()

                    logger.info(
                        f"File validated and saved: {filename}, size: {data['file_size']}, hash: {data['file_hash'][:16]}..."
                    )

        # Create evidence package
        evidence_package = evidence_workflow.processor.create_evidence_package(data)

        # Save to database (using Analysis model for now)
        analysis = Analysis(
            user_id=current_user.id,
            filename=data.get("filename", "evidence"),
            file_hash=data.get("file_hash"),
            file_size=data.get("file_size"),
            file_path=data.get("file_path"),
            status="uploaded",
            case_number=data["case_number"],
            evidence_number=data.get("id"),
            acquired_by=data["acquired_by"],
            source=data["source"],
        )
        analysis.generate_id()

        db.session.add(analysis)
        db.session.commit()

        # Save evidence package metadata
        metadata_path = (
            Path(app.config.get("ANALYSIS_FOLDER", "./bwc_analysis"))
            / analysis.id
            / "evidence_package.json"
        )
        metadata_path.parent.mkdir(parents=True, exist_ok=True)
        with open(metadata_path, "w") as f:
            json.dump(evidence_package, f, indent=2)

        return success_response(
            data={"evidence_id": analysis.id, "case_number": data["case_number"]},
            message="Evidence submitted successfully",
            status_code=201,
        )

    except Exception as e:
        # Log full error server-side
        logger.error(f"Evidence intake failed: {type(e).__name__}: {e}", exc_info=True)

        # Generate error ticket for support
        error_ticket = ErrorSanitizer.create_error_ticket()

        # Return sanitized error to user
        user_message = ErrorSanitizer.sanitize_error(e, "file")
        return error_response(
            user_message, error_code="OPERATION_FAILED", status_code=500, error_ticket=error_ticket
        )


@app.route("/api/evidence/list", methods=["GET"])
@login_required
def list_evidence():
    """List all evidence items for current user with pagination"""
    # Add pagination
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 20, type=int)
    per_page = min(per_page, 100)  # Cap at 100

    analyses_query = Analysis.query.filter_by(user_id=current_user.id).order_by(
        Analysis.created_at.desc()
    )
    analyses = analyses_query.limit(per_page).offset((page - 1) * per_page).all()

    evidence_list = []
    for analysis in analyses:
        # Load evidence package if exists
        metadata_path = (
            Path(app.config.get("ANALYSIS_FOLDER", "./bwc_analysis"))
            / analysis.id
            / "evidence_package.json"
        )

        if metadata_path.exists():
            with open(metadata_path, "r") as f:
                evidence_package = json.load(f)
            evidence_list.append(evidence_package)
        else:
            # Fallback to basic analysis data
            evidence_list.append(
                {
                    "evidence_id": analysis.id,
                    "case_information": {"case_number": analysis.case_number or "N/A"},
                    "evidence_details": {"description": analysis.filename, "type": "bwc_video"},
                    "processing_status": {
                        "stage": analysis.status,
                        "priority": "normal",
                        "sla_deadline": (datetime.utcnow() + timedelta(hours=72)).isoformat(),
                    },
                }
            )

    return jsonify(evidence_list)


# Analytics & AI Routes
@app.route("/analytics")
@login_required
def analytics_dashboard():
    """Analytics dashboard with charts and KPIs"""
    return send_file("templates/analytics.html")


@app.route("/api/ai/suggest", methods=["POST"])
@login_required
def ai_suggest():
    """AI-powered suggestions for forms"""
    from ai_suggestions import smart_suggest

    data = request.get_json()
    field = data.get("field")
    text = data.get("text", "")
    context = data.get("context", {})

    if field == "description":
        suggestions = smart_suggest.suggest_description(text)
        return jsonify({"suggestions": suggestions})

    elif field == "priority":
        result = smart_suggest.suggest_priority(text)
        return jsonify(result)

    elif field == "tags":
        evidence_type = context.get("evidence_type")
        result = smart_suggest.suggest_tags(text, evidence_type)
        return jsonify(result)

    elif field == "categorize":
        filename = context.get("filename", "")
        result = smart_suggest.auto_categorize(filename, text)
        return jsonify(result)

    elif field == "case_number":
        result = smart_suggest.suggest_case_number()
        return jsonify(result)

    elif field == "similar_cases":
        results = smart_suggest.suggest_similar_cases(text)
        return jsonify({"cases": results})

    elif field == "processing_time":
        priority = context.get("priority", "normal")
        evidence_type = context.get("evidence_type", "unknown")
        result = smart_suggest.predict_processing_time(priority, evidence_type, text)
        return jsonify(result)

    else:
        return jsonify({"error": "Unknown field"}), 400


# Legal AI Agents Routes
@app.route("/agents")
@login_required
def agents_page():
    """AI agents deployment and management page"""
    return render_template("agents.html", user=current_user)


@app.route("/command-center")
@login_required
def command_center():
    """AI Command Center - Main hub for evidence analysis and workflow"""
    return render_template("command-center.html", user=current_user)


# ========================================
# LEGAL ANALYSIS TOOLS
# ========================================

# Import legal analysis tools
try:
    from case_law_violation_scanner import ViolationScanner
    from statutory_compliance_checker import StatutoryComplianceChecker

    LEGAL_TOOLS_AVAILABLE = True
except ImportError:
    LEGAL_TOOLS_AVAILABLE = False
    print("[WARN] Legal analysis tools not available")


@app.route("/legal-analysis")
@login_required
def legal_analysis_dashboard():
    """Legal analysis dashboard with violation scanner and compliance checker"""
    return render_template("legal-analysis.html", user=current_user)


@app.route("/api/legal/scan-violations", methods=["POST"])
@login_required
@require_tier(TierLevel.PROFESSIONAL)
@check_usage_limit("legal_analyses_per_month", increment=1)
def scan_violations():
    """Scan transcript for legal violations"""
    if not LEGAL_TOOLS_AVAILABLE:
        return jsonify({"error": "Legal analysis tools not available"}), 503

    try:
        data = request.get_json()
        transcript = data.get("transcript", "")
        context = data.get("context", {})

        if not transcript:
            return jsonify({"error": "No transcript provided"}), 400

        scanner = ViolationScanner()
        results = scanner.scan_transcript(transcript, context)

        return jsonify(results), 200

    except Exception as e:
        app.logger.error(f"Violation scan error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/legal/check-compliance", methods=["POST"])
@login_required
@require_tier(TierLevel.PROFESSIONAL)
@check_usage_limit("legal_analyses_per_month", increment=1)
def check_compliance():
    """Check evidence for statutory compliance"""
    if not LEGAL_TOOLS_AVAILABLE:
        return jsonify({"error": "Legal analysis tools not available"}), 503

    try:
        data = request.get_json()
        evidence = data.get("evidence", {})

        if not evidence:
            return jsonify({"error": "No evidence provided"}), 400

        checker = StatutoryComplianceChecker()
        results = checker.comprehensive_check(evidence)

        return jsonify(results), 200

    except Exception as e:
        app.logger.error(f"Compliance check error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/legal/combined-analysis", methods=["POST"])
@login_required
@require_tier(TierLevel.PROFESSIONAL)
@check_usage_limit("legal_analyses_per_month", increment=1)
def combined_legal_analysis():
    """Run both violation scan and compliance check"""
    if not LEGAL_TOOLS_AVAILABLE:
        return jsonify({"error": "Legal analysis tools not available"}), 503

    try:
        data = request.get_json()
        transcript = data.get("transcript", "")
        evidence = data.get("evidence", {})

        scanner = ViolationScanner()
        checker = StatutoryComplianceChecker()

        violation_results = scanner.scan_transcript(transcript, evidence)
        compliance_results = checker.comprehensive_check(evidence)

        # Combine results
        combined = {
            "analysis_timestamp": datetime.utcnow().isoformat(),
            "evidence_id": evidence.get("id", "Unknown"),
            "violations": violation_results,
            "compliance": compliance_results,
            "overall_assessment": {
                "total_issues": violation_results["total_violations"]
                + compliance_results["total_issues"],
                "critical_violations": len(violation_results.get("critical_violations", [])),
                "non_compliant_count": compliance_results["issues_by_status"]["non_compliant"],
                "recommended_actions": violation_results.get("recommended_motions", [])
                + compliance_results.get("recommendations", []),
            },
        }

        return jsonify(combined), 200

    except Exception as e:
        app.logger.error(f"Combined analysis error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


# ========================================
# PHASE 1 PREMIUM FEATURES - API ROUTES
# ========================================


@app.route("/api/evidence/transcribe", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("transcription_minutes_per_month", increment=1)
def transcribe_audio():
    """Transcribe audio/video using Whisper AI"""
    if not WHISPER_AVAILABLE:
        return jsonify({"error": "Whisper transcription service not available"}), 503

    try:
        # Get file from request
        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No file selected"}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(file_path)

        # Initialize Whisper service
        whisper_service = WhisperTranscriptionService(model_size="base")

        # Get options from request
        language = request.form.get("language", None)
        enable_timestamps = request.form.get("timestamps", "true").lower() == "true"

        # Transcribe
        result = whisper_service.transcribe_audio(
            str(file_path), language=language, enable_timestamps=enable_timestamps
        )

        # Export to requested format
        export_format = request.form.get("format", "json")
        if export_format == "json":
            return jsonify({"success": True, "transcription": result})
        else:
            output = whisper_service.export_transcript(result, export_format)
            return output, 200, {"Content-Type": "text/plain"}

    except Exception as e:
        app.logger.error(f"Transcription error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/evidence/ocr", methods=["POST"])
@login_required
@require_tier(TierLevel.FREE)
@check_usage_limit("document_pages_per_month", increment=1)
def extract_text_ocr():
    """Extract text from images/scanned PDFs using OCR"""
    if not OCR_AVAILABLE:
        return jsonify({"error": "OCR service not available"}), 503

    try:
        # Get file from request
        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No file selected"}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(file_path)

        # Initialize OCR service
        engine = request.form.get("engine", "tesseract")  # or 'aws'
        ocr_service = OCRService(engine=engine)

        # Get options
        language = request.form.get("language", "eng")
        preserve_layout = request.form.get("preserve_layout", "false").lower() == "true"

        # Detect file type and process
        file_ext = os.path.splitext(filename)[1].lower()

        if file_ext == ".pdf":
            result = ocr_service.extract_text_from_pdf(str(file_path), language)
        else:
            result = ocr_service.extract_text_from_image(str(file_path), language, preserve_layout)

        return jsonify({"success": True, "ocr_result": result})

    except Exception as e:
        app.logger.error(f"OCR error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/account/2fa")
@login_required
def two_factor_setup_page():
    """2FA setup page"""
    if not TWO_FACTOR_AVAILABLE:
        flash("Two-factor authentication is not available", "warning")
        return redirect(url_for("account_settings"))

    return render_template("auth/2fa-setup.html")


@app.route("/api/account/2fa/setup", methods=["POST"])
@login_required
def setup_two_factor():
    """Setup 2FA for current user"""
    if not TWO_FACTOR_AVAILABLE:
        return jsonify({"error": "2FA service not available"}), 503

    try:
        tfa_service = TwoFactorAuthService(issuer_name="BarberX Legal")

        # Generate 2FA setup data
        setup_data = tfa_service.setup_2fa_for_user(current_user.email)

        # Store secret in database (encrypted in production!)
        # For now, return to frontend - implement User2FA model later

        return jsonify(
            {
                "success": True,
                "qr_code": setup_data["qr_code"],
                "manual_entry_key": setup_data["manual_entry_key"],
                "backup_codes": setup_data["backup_codes"],
            }
        )

    except Exception as e:
        app.logger.error(f"2FA setup error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/account/2fa/verify", methods=["POST"])
@login_required
def verify_two_factor():
    """Verify 2FA token"""
    if not TWO_FACTOR_AVAILABLE:
        return jsonify({"error": "2FA service not available"}), 503

    try:
        data = request.get_json()
        token = data.get("token")
        secret = data.get("secret")  # In production, fetch from database

        if not token or not secret:
            return jsonify({"error": "Missing token or secret"}), 400

        tfa_service = TwoFactorAuthService()
        is_valid = tfa_service.verify_token(secret, token)

        if is_valid:
            # Enable 2FA for user in database
            return jsonify({"success": True, "message": "2FA enabled successfully"})
        else:
            return jsonify({"success": False, "message": "Invalid token"}), 401

    except Exception as e:
        app.logger.error(f"2FA verification error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/pricing")
def pricing_page():
    """Pricing page with subscription plans"""
    if STRIPE_AVAILABLE:
        plans = SUBSCRIPTION_PLANS
    else:
        # Fallback plans if Stripe not available
        plans = {
            "basic": {"name": "Basic", "price_monthly": 29},
            "pro": {"name": "Professional", "price_monthly": 99},
            "premium": {"name": "Premium", "price_monthly": 299},
            "enterprise": {"name": "Enterprise", "price_monthly": "custom"},
        }

    return render_template("pricing.html", plans=plans)


@app.route("/api/billing/create-checkout", methods=["POST"])
@login_required
def create_checkout_session():
    """Create Stripe checkout session"""
    if not STRIPE_AVAILABLE:
        return jsonify({"error": "Payment processing not available"}), 503

    try:
        data = request.get_json()
        plan = data.get("plan")  # 'basic', 'pro', 'premium'
        billing_period = data.get("period", "monthly")  # 'monthly' or 'yearly'

        if plan not in SUBSCRIPTION_PLANS:
            return jsonify({"error": "Invalid plan"}), 400

        # Initialize Stripe service
        stripe_service = StripePaymentService()

        # Get or create Stripe customer
        stripe_customer_id = getattr(current_user, "stripe_customer_id", None)

        if not stripe_customer_id:
            customer = stripe_service.create_customer(
                email=current_user.email,
                name=current_user.username,
                metadata={"user_id": current_user.id},
            )
            stripe_customer_id = customer["customer_id"]

            # Save to database
            current_user.stripe_customer_id = stripe_customer_id
            db.session.commit()

        # Get price ID
        price_key = f"stripe_price_id_{billing_period}"
        price_id = SUBSCRIPTION_PLANS[plan].get(price_key)

        if not price_id:
            return jsonify({"error": "Price ID not configured"}), 500

        # Create checkout session
        success_url = url_for("billing_success", _external=True)
        cancel_url = url_for("pricing_page", _external=True)

        session = stripe_service.create_checkout_session(
            customer_id=stripe_customer_id,
            price_id=price_id,
            success_url=success_url,
            cancel_url=cancel_url,
            trial_days=14,
        )

        return jsonify(
            {"success": True, "checkout_url": session["url"], "session_id": session["session_id"]}
        )

    except Exception as e:
        app.logger.error(f"Checkout creation error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/billing/webhook", methods=["POST"])
def stripe_webhook():
    """Handle Stripe webhook events"""
    if not STRIPE_AVAILABLE:
        return jsonify({"error": "Payment processing not available"}), 503

    try:
        payload = request.data
        sig_header = request.headers.get("Stripe-Signature")

        stripe_service = StripePaymentService()
        event_data = stripe_service.handle_webhook(payload, sig_header)

        # Handle different event types
        action = event_data.get("action")

        if action == "subscription_created":
            # Update user subscription status
            app.logger.info(f"Subscription created: {event_data}")
        elif action == "subscription_deleted":
            # Handle cancellation
            app.logger.info(f"Subscription deleted: {event_data}")
        elif action == "payment_failed":
            # Handle failed payment
            app.logger.warning(f"Payment failed: {event_data}")

        return jsonify({"received": True})

    except Exception as e:
        app.logger.error(f"Webhook error: {str(e)}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/billing/success")
@login_required
def billing_success():
    """Billing success page"""
    flash("Subscription activated successfully! Welcome to BarberX Premium.", "success")
    return redirect(url_for("evidence_dashboard"))


@app.route("/account/billing")
@login_required
def billing_dashboard():
    """User billing dashboard"""
    return render_template("auth/billing.html")


@app.route("/offline.html")
def offline_page():
    """PWA offline fallback page"""
    return render_template("offline.html")


# PWA static files
@app.route("/manifest.json")
def pwa_manifest():
    """Serve PWA manifest"""
    return send_file("manifest.json", mimetype="application/json")


@app.route("/service-worker.js")
def service_worker():
    """Serve service worker"""
    return send_file("service-worker.js", mimetype="application/javascript")


# ========================================
# END PHASE 1 ROUTES
# ========================================


@app.route("/integrated-analysis")
@login_required
def integrated_analysis_page():
    """Unified evidence analysis with chat and document generation"""
    return render_template("integrated-analysis.html", user=current_user)


@app.route("/api/agents/deploy", methods=["POST"])
@login_required
def deploy_agent():
    """Deploy a new AI agent"""
    from legal_ai_agents import agent_manager

    data = request.get_json()
    agent_type = data.get("agent_type")
    config = data.get("config", {})

    try:
        agent_id = agent_manager.deploy_agent(
            agent_type=agent_type, user_id=str(current_user.id), config=config
        )

        return jsonify({"agent_id": agent_id, "message": "Agent deployed successfully"})
    except Exception as e:
        app.logger.error(f"Agent deployment error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/api/agents/list", methods=["GET"])
@login_required
def list_agents():
    """List all agents for current user"""
    from legal_ai_agents import agent_manager

    try:
        agents = agent_manager.list_user_agents(str(current_user.id))
        return jsonify(agents)
    except Exception as e:
        app.logger.error(f"Agent list error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/api/agents/execute/<agent_id>", methods=["POST"])
@login_required
def execute_agent(agent_id):
    """Execute an agent with input data"""
    from legal_ai_agents import agent_manager

    data = request.get_json()
    input_data = data.get("input_data", {})

    try:
        result = agent_manager.execute_agent(agent_id, input_data)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f"Agent execution error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/api/evidence/analyze-pdf", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("pdf_documents_per_month", increment=1)
def analyze_pdf_discovery():
    """Analyze PDF discovery document for legal information"""
    from enhanced_pdf_discovery_analyzer import PDFDiscoveryAnalyzer

    try:
        data = request.get_json()
        pdf_content = data.get("content", "")
        filename = data.get("filename", "document.pdf")

        if not pdf_content:
            return jsonify({"error": "No content provided"}), 400

        # Analyze the PDF
        analyzer = PDFDiscoveryAnalyzer()
        results = analyzer.analyze_document(pdf_content, filename)

        # Save to analysis record if user wants
        if data.get("save_analysis"):
            # TODO: Save to database
            pass

        return jsonify(
            {
                "success": True,
                "results": results,
                "formatted_report": analyzer.export_to_report(results),
            }
        )

    except Exception as e:
        app.logger.error(f"PDF analysis error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "default"),
            error_code="OPERATION_FAILED",
            status_code=500,
            error_ticket=error_ticket,
        )


@app.route("/api/agents/status/<agent_id>", methods=["GET"])
@login_required
def get_agent_status(agent_id):
    """Get agent status and results"""
    from legal_ai_agents import agent_manager

    try:
        status = agent_manager.get_agent_status(agent_id)
        return jsonify(status)
    except Exception as e:
        app.logger.error(f"Agent status error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "database"),
            error_code="OPERATION_FAILED",
            status_code=404,
            error_ticket=error_ticket,
        )


@app.route("/api/agents/<agent_id>", methods=["DELETE"])
@login_required
def delete_agent(agent_id):
    """Delete an agent"""
    from legal_ai_agents import agent_manager

    try:
        agent_manager.delete_agent(agent_id)
        return jsonify({"message": "Agent deleted"})
    except Exception as e:
        app.logger.error(f"Agent deletion error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


# ============================================================================
# UNIFIED INTEGRATION ROUTES - Connect all features into one workflow
# ============================================================================


@app.route("/api/workflow/process-evidence", methods=["POST"])
@login_required
def process_evidence_workflow():
    """
    Complete workflow: Upload → Auto-analyze → Deploy agents → Enable chat
    """
    from unified_integration import get_orchestrator

    try:
        data = request.get_json()
        orchestrator = get_orchestrator(current_user.id)

        result = orchestrator.process_evidence_intake(data)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f"Evidence workflow error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/api/workflow/chat", methods=["POST"])
@login_required
def workflow_chat():
    """AI chat with full evidence context"""
    from unified_integration import get_orchestrator

    try:
        data = request.get_json()
        workflow_id = data.get("workflow_id")
        query = data.get("query", "")

        orchestrator = get_orchestrator(current_user.id)
        result = orchestrator.process_ai_chat_query(workflow_id, query)

        return jsonify(result)
    except Exception as e:
        app.logger.error(f"Workflow chat error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/api/workflow/generate-document", methods=["POST"])
@login_required
def workflow_generate_document():
    """Generate legal document from evidence analysis"""
    from unified_integration import get_orchestrator

    try:
        data = request.get_json()
        workflow_id = data.get("workflow_id")
        document_type = data.get("document_type")
        custom_inputs = data.get("custom_inputs", {})

        orchestrator = get_orchestrator(current_user.id)
        result = orchestrator.generate_document_from_analysis(
            workflow_id, document_type, custom_inputs
        )

        return jsonify(result)
    except Exception as e:
        app.logger.error(f"Document generation error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/api/workflow/scan-document", methods=["POST"])
@login_required
def workflow_scan_document():
    """Scan and process document with OCR + AI"""
    from unified_integration import get_orchestrator

    try:
        # Handle file upload
        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "Empty filename"}), 400

        # Read file data (fix: only read once)
        file_content = file.read()
        file_data = {
            "name": secure_filename(file.filename),
            "content": file_content.decode("utf-8", errors="ignore"),
            "size": len(file_content),
        }

        orchestrator = get_orchestrator(current_user.id)
        result = orchestrator.scan_and_process_document(file_data)

        return jsonify(result)
    except Exception as e:
        app.logger.error(f"Document scan error: {e}")
        # Log error server-side
        logger.error(f"{request.path} failed: {type(e).__name__}: {e}", exc_info=True)
        error_ticket = ErrorSanitizer.create_error_ticket()
        return error_response(
            ErrorSanitizer.sanitize_error(e, "validation"),
            error_code="OPERATION_FAILED",
            status_code=400,
            error_ticket=error_ticket,
        )


@app.route("/tools")
@login_required
def tools_index():
    """Tools hub - unified AI analysis tools interface"""
    from models_auth import UsageTracking
    
    try:
        usage = UsageTracking.get_or_create_current(current_user.id)
        limits = current_user.get_tier_limits()
        return render_template("tools-hub.html", user=current_user, usage=usage, limits=limits)
    except Exception as e:
        app.logger.error(f"Tools hub error: {e}")
        # Fallback with minimal data
        usage = type('obj', (object,), {
            'bwc_videos_processed': 0,
            'pdf_documents_processed': 0,
            'transcription_minutes_used': 0,
            'storage_used_mb': 0
        })()
        return render_template("tools-hub.html", user=current_user, usage=usage, limits={})


@app.route("/tools/transcript")
@login_required
def transcript_search():
    """Transcript search tool"""
    return render_template("tools/transcript.html")


@app.route("/tools/entity-extract")
@login_required
def entity_extract():
    """Entity extraction tool"""
    return render_template("tools/entity-extract.html")


@app.route("/tools/timeline")
@login_required
def timeline_builder():
    """Timeline builder tool"""
    return render_template("tools/timeline.html")


@app.route("/tools/discrepancy")
@login_required
def discrepancy_finder():
    """Discrepancy finder tool"""
    return render_template("tools/discrepancy.html")


@app.route("/tools/batch")
@login_required
def batch_processor():
    """Batch processor tool"""
    return render_template("tools/batch.html")


@app.route("/tools/api")
@login_required
def api_console():
    """API console tool"""
    return render_template("tools/api-console.html")


@app.route("/tools/ocr")
@login_required
def ocr_tool():
    """OCR text extraction tool"""
    return render_template("tools/ocr.html")


@app.route("/tools/case-law")
@login_required
def case_law_tool():
    """Case law research tool"""
    return render_template("tools/case-law.html")


@app.route("/batch-upload")
@login_required
def batch_upload_page():
    """Unified batch upload page for PDFs, BWC videos, and images"""
    return send_file("templates/batch-upload-unified.html")


# ========================================
# RESOURCE PAGES
# ========================================


@app.route("/docs")
def docs():
    """Documentation home"""
    return send_file("templates/resources/docs.html")


@app.route("/api")
def api_reference():
    """API reference"""
    return send_file("templates/resources/api-reference.html")


@app.route("/blog")
def blog():
    """Blog home"""
    return send_file("templates/resources/blog.html")


@app.route("/case-studies")
def case_studies():
    """Case studies"""
    return send_file("templates/resources/case-studies.html")


@app.route("/guides")
def guides():
    """User guides"""
    return send_file("templates/resources/guides.html")


@app.route("/faq")
def faq():
    """FAQ page"""
    return send_file("templates/resources/faq.html")


@app.route("/animation-demo")
@app.route("/animation-demo.html")
def animation_demo():
    """Animation system demo page"""
    return send_file("animation-demo.html")


# ========================================
# COMPANY PAGES
# ========================================


@app.route("/about")
def about():
    """About us"""
    return send_file("templates/company/about.html")


@app.route("/careers")
def careers():
    """Careers page"""
    return send_file("templates/company/careers.html")


@app.route("/licenses")
@app.route("/company/licenses")
def licenses():
    """Open source licenses and attributions"""
    return send_file("templates/company/licenses.html")


@app.route("/contact")
def contact():
    """Contact page"""
    return send_file("templates/company/contact.html")


@app.route("/honor")
def military_honor_installation():
    """
    US Military Honor Installation

    Dedicated to all who served, are serving, and made the ultimate sacrifice.
    Displays official US military flags in proper protocol and order per Title 4 USC.
    Honors religious martyrs who died for freedom of conscience.

    'By the Grace of Almighty God, we honor their service and defend the Constitution they protected.'
    """
    return render_template("honor.html")


@app.route("/founding-documents")
def founding_documents():
    """
    Complete archive of US founding documents

    The Constitution for the United States of America, Bill of Rights, all amendments,
    Declaration of Independence, state constitutions, and foundational legal texts.
    Official government sources with full text for legal reference.

    'Truth is the only subject' - The supreme law of the real land of North America.
    """
    return render_template("founding-documents.html")


@app.route("/components/header-witness")
def header_witness_component():
    """
    Constitutional Header Witness Component

    'Let a Header Be Set as Witness.'

    Establishes a unified mark at the summit of the site, constituted as a quiet
    standard and not an ornament. The US Flag stands first by order, proportion,
    and Title 4 USC authority. All subordinate standards revealed by consent or inquiry.

    Set the Standard first. Reveal the rest by order. Alter nothing. Animate nothing.
    Bind by hierarchy and truth. Let the header stand as witness.
    """
    return render_template("components/header-witness.html")


@app.route("/press")
def press():
    """Press/media page"""
    return send_file("templates/company/press.html")


# ========================================
# STATIC ASSETS
# ========================================


@app.route("/assets/<path:filename>")
def serve_assets(filename):
    """Serve static assets from assets folder"""
    return send_file(os.path.join("assets", filename))


# ========================================
# API ROUTES - Public/Unauthenticated
# ========================================


@app.route("/api/founding-member-signup", methods=["POST", "OPTIONS"])
@csrf.exempt
def founding_member_signup():
    """
    Handle Founding Member email capture

    Stores email, name, firm in database for Founding Member program.
    First 100 members get lifetime $19/month rate lock.

    Returns:
        - success: true/false
        - spots_remaining: int (out of 100)
        - message: str
    """
    # Handle OPTIONS preflight
    if request.method == "OPTIONS":
        return jsonify({"status": "ok"}), 200

    try:
        data = request.get_json()

        if not data or "email" not in data:
            return jsonify({"success": False, "message": "Email is required"}), 400

        email = data.get("email", "").strip().lower()
        name = data.get("name", "").strip()
        firm = data.get("firm", "").strip()
        source = data.get("source", "unknown")

        # Basic email validation
        if "@" not in email or "." not in email:
            return jsonify({"success": False, "message": "Invalid email address"}), 400

        # Use simple flat file for email capture (avoid database complexity)
        # Production version should integrate with Stripe customer creation
        import csv
        import os
        from datetime import datetime

        signups_file = Path("founding_member_signups.csv")

        # Check if already signed up
        if signups_file.exists():
            with open(signups_file, "r", newline="", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    if row.get("email", "").lower() == email:
                        # Count total signups
                        f.seek(0)
                        count = sum(1 for _ in reader) - 1  # -1 for header
                        spots_remaining = max(0, 100 - count)
                        return (
                            jsonify(
                                {
                                    "success": True,
                                    "message": "You're already on the list!",
                                    "spots_remaining": spots_remaining,
                                }
                            ),
                            200,
                        )

        # Append new signup
        file_exists = signups_file.exists()
        with open(signups_file, "a", newline="", encoding="utf-8") as f:
            fieldnames = ["email", "name", "firm", "source", "signup_date", "status"]
            writer = csv.DictWriter(f, fieldnames=fieldnames)

            if not file_exists:
                writer.writeheader()

            writer.writerow(
                {
                    "email": email,
                    "name": name,
                    "firm": firm,
                    "source": source,
                    "signup_date": datetime.utcnow().isoformat(),
                    "status": "pending",
                }
            )

        # Count total signups
        with open(signups_file, "r", newline="", encoding="utf-8") as f:
            count = sum(1 for _ in f) - 1  # -1 for header
            spots_remaining = max(0, 100 - count)

        app.logger.info(
            f"Founding Member signup: {email} ({name}) from {firm} - {spots_remaining} spots remaining"
        )

        # TODO: Send welcome email with checkout link
        # TODO: Integrate with Stripe checkout session

        return (
            jsonify(
                {
                    "success": True,
                    "message": "Success! Check your email for next steps.",
                    "spots_remaining": spots_remaining,
                }
            ),
            200,
        )

    except Exception as e:
        app.logger.error(f"Founding Member signup error: {e}")
        return (
            jsonify(
                {
                    "success": False,
                    "message": "An error occurred. Please try again or contact founders@barberx.info",
                }
            ),
            500,
        )


# ========================================
# API ROUTES - Analysis
# ========================================


@app.route("/api/upload", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("bwc_videos_per_month", increment=1)
def upload_file():
    """Handle BWC video file upload"""

    # Check if user can analyze
    if not current_user.can_analyze():
        return jsonify({"error": "Monthly analysis limit reached. Please upgrade your plan."}), 403

    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]

    # Validate file using security helper
    is_valid, error_msg = validate_upload_file(file, ALLOWED_VIDEO_EXTENSIONS)
    if not is_valid:
        return jsonify({"error": error_msg}), 400

    # Secure filename
    filename = secure_filename(file.filename)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{current_user.id}_{timestamp}_{filename}"
    filepath = app.config["UPLOAD_FOLDER"] / unique_filename

    # Save file
    file.save(filepath)

    # Calculate hash
    file_hash = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            file_hash.update(chunk)
    file_hash_hex = file_hash.hexdigest()

    file_size = os.path.getsize(filepath)

    # Create analysis record
    analysis = Analysis(
        user_id=current_user.id,
        filename=filename,
        file_hash=file_hash_hex,
        file_size=file_size,
        file_path=str(filepath),
        status="uploaded",
    )
    analysis.generate_id()

    db.session.add(analysis)
    db.session.commit()

    # Update user storage
    current_user.storage_used_mb += file_size / (1024 * 1024)
    db.session.commit()

    # Log audit
    AuditLog.log("file_uploaded", "analysis", analysis.id, {"filename": filename})

    app.logger.info(f"File uploaded: {filename} by user {current_user.email}")

    return jsonify(
        {
            "upload_id": analysis.id,
            "filename": filename,
            "file_hash": file_hash_hex,
            "file_size": file_size,
            "message": "File uploaded successfully",
        }
    )


@app.route("/api/upload/pdf", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("pdf_documents_per_month", increment=1)
def upload_pdf():
    """Handle single PDF file upload"""
    import time

    from usage_meter import SmartMeter, UsageQuota

    # Check file quota
    quota = UsageQuota.query.filter_by(user_id=current_user.id).first()
    if not quota:
        quota = SmartMeter.initialize_user_quota(current_user.id)

    has_quota, error_msg = quota.check_quota("files")
    if not has_quota:
        SmartMeter.track_event(
            event_type="upload_blocked",
            event_category="quota",
            status="denied",
            error_message=error_msg,
        )
        return (
            jsonify(
                {"error": "Upload quota exceeded", "message": error_msg, "upgrade_url": "/pricing"}
            ),
            429,
        )

    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    # Validate PDF file
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF files are allowed"}), 400

    # Get optional metadata
    case_number = request.form.get("case_number", "")
    document_type = request.form.get("document_type", "")
    description = request.form.get("description", "")
    tags_str = request.form.get("tags", "")
    tags = [t.strip() for t in tags_str.split(",") if t.strip()] if tags_str else []

    # Secure filename
    original_filename = file.filename
    filename = secure_filename(file.filename)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{timestamp}_{filename}"

    # Create upload directory
    upload_dir = Path("./uploads/pdfs")
    upload_dir.mkdir(parents=True, exist_ok=True)

    filepath = upload_dir / unique_filename

    start_time = time.time()
    # Save file
    file.save(filepath)

    file_size = filepath.stat().st_size
    duration = time.time() - start_time

    # Track upload
    SmartMeter.track_event(
        event_type="file_upload",
        event_category="storage",
        resource_name=original_filename,
        quantity=1,
        file_size_bytes=file_size,
        duration_seconds=duration,
        status="success",
    )

    # Update storage quota
    quota.increment_quota("files")
    quota.increment_quota("storage", file_size)
    db.session.commit()

    # Get file size
    file_size = os.path.getsize(filepath)

    # Create PDF upload record
    pdf_upload = PDFUpload(
        user_id=current_user.id if current_user.is_authenticated else None,
        filename=unique_filename,
        original_filename=original_filename,
        file_path=str(filepath),
        file_size=file_size,
        case_number=case_number,
        document_type=document_type,
        description=description,
        tags=tags,
        status="uploaded",
    )

    # Generate hash
    pdf_upload.generate_hash(str(filepath))

    db.session.add(pdf_upload)
    db.session.commit()

    # Log audit
    AuditLog.log(
        "pdf_uploaded",
        "pdf_upload",
        str(pdf_upload.id),
        {"filename": original_filename, "file_size": file_size},
    )

    app.logger.info(f"PDF uploaded: {original_filename} (ID: {pdf_upload.id})")

    return jsonify(
        {
            "success": True,
            "upload_id": pdf_upload.id,
            "filename": original_filename,
            "file_hash": pdf_upload.file_hash,
            "file_size": file_size,
            "message": "PDF uploaded successfully",
        }
    )


@app.route("/api/upload/pdf/batch", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("pdf_documents_per_month", increment=1)
def batch_upload_pdf():
    """Handle batch PDF file upload"""
    if "files" not in request.files:
        return jsonify({"error": "No files provided"}), 400

    files = request.files.getlist("files")

    if not files:
        return jsonify({"error": "No files selected"}), 400

    results = {"total": len(files), "successful": [], "failed": []}

    for file in files:
        try:
            if file.filename == "":
                results["failed"].append({"filename": "unknown", "error": "Empty filename"})
                continue

            # Validate PDF
            if not file.filename.lower().endswith(".pdf"):
                results["failed"].append({"filename": file.filename, "error": "Not a PDF file"})
                continue

            # Secure filename
            original_filename = file.filename
            filename = secure_filename(file.filename)
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S_%f")
            unique_filename = f"{timestamp}_{filename}"

            # Create upload directory
            upload_dir = Path("./uploads/pdfs")
            upload_dir.mkdir(parents=True, exist_ok=True)

            filepath = upload_dir / unique_filename

            # Save file
            file.save(filepath)

            # Get file size
            file_size = os.path.getsize(filepath)

            # Create PDF upload record
            pdf_upload = PDFUpload(
                user_id=current_user.id if current_user.is_authenticated else None,
                filename=unique_filename,
                original_filename=original_filename,
                file_path=str(filepath),
                file_size=file_size,
                status="uploaded",
            )

            # Generate hash
            pdf_upload.generate_hash(str(filepath))

            db.session.add(pdf_upload)
            db.session.commit()

            results["successful"].append(
                {
                    "upload_id": pdf_upload.id,
                    "filename": original_filename,
                    "file_size": file_size,
                    "file_hash": pdf_upload.file_hash,
                }
            )

            app.logger.info(f"PDF uploaded (batch): {original_filename} (ID: {pdf_upload.id})")

        except Exception as e:
            results["failed"].append(
                {"filename": file.filename if file else "unknown", "error": str(e)}
            )
            app.logger.error(f"Batch PDF upload error: {str(e)}")

    # Log audit for batch
    AuditLog.log(
        "pdf_batch_uploaded",
        "pdf_upload",
        None,
        {
            "total": results["total"],
            "successful": len(results["successful"]),
            "failed": len(results["failed"]),
        },
    )

    return jsonify(
        {
            "success": True,
            "results": results,
            "summary": {
                "total": results["total"],
                "successful": len(results["successful"]),
                "failed": len(results["failed"]),
            },
        }
    )


@app.route("/api/pdfs", methods=["GET"])
@login_required
def list_pdfs():
    """List uploaded PDFs"""
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 50, type=int)

    # Filter by user if authenticated
    query = PDFUpload.query

    if current_user.is_authenticated:
        # Admins can see all, users see only their own
        if current_user.role != "admin":
            query = query.filter_by(user_id=current_user.id)
    else:
        # Public access - only show public PDFs
        query = query.filter_by(is_public=True)

    pagination = query.order_by(PDFUpload.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )

    return jsonify(
        {
            "total": pagination.total,
            "page": page,
            "per_page": per_page,
            "total_pages": pagination.pages,
            "pdfs": [pdf.to_dict() for pdf in pagination.items],
        }
    )


@app.route("/api/pdf/<int:pdf_id>", methods=["GET"])
@login_required
def get_pdf_info(pdf_id):
    """Get PDF information"""
    pdf = PDFUpload.query.get(pdf_id)

    if not pdf:
        return jsonify({"error": "PDF not found"}), 404

    # Check access
    if not pdf.is_public and (
        not current_user.is_authenticated
        or (current_user.id != pdf.user_id and current_user.role != "admin")
    ):
        return jsonify({"error": "Access denied"}), 403

    return jsonify(pdf.to_dict())


@app.route("/api/pdf/<int:pdf_id>/download", methods=["GET"])
@login_required
def download_pdf(pdf_id):
    """Download PDF file"""
    pdf = PDFUpload.query.get(pdf_id)

    if not pdf:
        return jsonify({"error": "PDF not found"}), 404

    # Check access
    if not pdf.is_public and (
        not current_user.is_authenticated
        or (current_user.id != pdf.user_id and current_user.role != "admin")
    ):
        return jsonify({"error": "Access denied"}), 403

    # Check if file exists
    if not os.path.exists(pdf.file_path):
        return jsonify({"error": "File not found on server"}), 404

    return send_file(
        pdf.file_path,
        as_attachment=True,
        download_name=pdf.original_filename,
        mimetype="application/pdf",
    )


@app.route("/api/pdf/<int:pdf_id>", methods=["DELETE"])
@login_required
def delete_pdf(pdf_id):
    """Delete PDF file"""
    pdf = PDFUpload.query.get(pdf_id)

    if not pdf:
        return jsonify({"error": "PDF not found"}), 404

    # Check permission
    if current_user.id != pdf.user_id and current_user.role != "admin":
        return jsonify({"error": "Permission denied"}), 403

    # Delete file from disk
    if os.path.exists(pdf.file_path):
        os.remove(pdf.file_path)

    # Delete from database
    db.session.delete(pdf)
    db.session.commit()

    # Log audit
    AuditLog.log("pdf_deleted", "pdf_upload", str(pdf_id), {"filename": pdf.original_filename})

    return jsonify({"success": True, "message": "PDF deleted successfully"})


@app.route("/api/analyze", methods=["POST"])
@login_required
def analyze_video():
    """Start BWC video analysis"""
    data = request.get_json()
    upload_id = data.get("upload_id")

    analysis = Analysis.query.filter_by(id=upload_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    # Update case metadata
    analysis.case_number = data.get("case_number", "")
    analysis.evidence_number = data.get("evidence_number", "")
    analysis.acquired_by = data.get("acquired_by", current_user.full_name)
    analysis.source = data.get("source", "Web Upload")
    analysis.known_officers = data.get("known_officers", [])
    analysis.status = "analyzing"
    analysis.started_at = datetime.utcnow()
    analysis.progress = 0

    db.session.commit()

    # Start analysis in background
    def run_analysis():
        try:
            analysis.current_step = "Extracting audio..."
            analysis.progress = 10
            db.session.commit()

            # Run forensic analysis
            report = analyzer.analyze_bwc_file(
                video_path=analysis.file_path,
                acquired_by=analysis.acquired_by,
                source=analysis.source,
                case_number=analysis.case_number,
                evidence_number=analysis.evidence_number,
                known_officers=analysis.known_officers,
            )

            analysis.progress = 90
            analysis.current_step = "Generating reports..."
            db.session.commit()

            # Save reports
            output_dir = app.config["ANALYSIS_FOLDER"] / analysis.id
            # Export report (output_files not used in mock mode)
            _ = analyzer.export_report(
                report, output_dir=str(output_dir), formats=["json", "txt", "md"]
            )

            # Update analysis record
            analysis.status = "completed"
            analysis.progress = 100
            analysis.current_step = "Analysis complete"
            analysis.completed_at = datetime.utcnow()
            analysis.duration = report.duration
            analysis.total_speakers = len(report.speakers)
            analysis.total_segments = len(report.transcript)
            analysis.total_discrepancies = len(report.discrepancies)
            analysis.critical_discrepancies = len(
                [d for d in report.discrepancies if d.severity == "critical"]
            )
            analysis.report_json_path = str(output_dir / "report.json")
            analysis.report_txt_path = str(output_dir / "report.txt")
            analysis.report_md_path = str(output_dir / "report.md")

            # Update user stats
            current_user.analyses_count += 1

            db.session.commit()

            # Log audit
            AuditLog.log("analysis_completed", "analysis", analysis.id)

            app.logger.info(f"Analysis completed: {analysis.id}")

        except Exception as e:
            analysis.status = "failed"
            analysis.error_message = str(e)
            analysis.progress = 0
            db.session.commit()

            app.logger.error(f"Analysis failed: {analysis.id} - {str(e)}")

    # Start background thread
    thread = threading.Thread(target=run_analysis)
    thread.start()

    analysis_tasks[upload_id] = thread

    return jsonify({"upload_id": upload_id, "message": "Analysis started", "status": "analyzing"})


@app.route("/api/analyses", methods=["GET"])
@login_required
def list_analyses():
    """List user's analyses"""
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 20, type=int)

    pagination = current_user.analyses.order_by(Analysis.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )

    return jsonify(
        {
            "total": pagination.total,
            "page": page,
            "per_page": per_page,
            "total_pages": pagination.pages,
            "analyses": [a.to_dict(include_results=True) for a in pagination.items],
        }
    )


@app.route("/api/analysis/<analysis_id>", methods=["GET"])
@login_required
def get_analysis(analysis_id):
    """Get analysis details"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    return jsonify(analysis.to_dict(include_results=True))


@app.route("/api/analysis/<analysis_id>/status", methods=["GET"])
@login_required
def get_analysis_status(analysis_id):
    """Get real-time analysis status"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    status_data = {
        "id": analysis.id,
        "status": analysis.status,
        "progress": analysis.progress,
        "current_step": analysis.current_step,
        "filename": analysis.filename,
        "case_number": analysis.case_number,
        "created_at": analysis.created_at.isoformat(),
        "started_at": analysis.started_at.isoformat() if analysis.started_at else None,
        "completed_at": analysis.completed_at.isoformat() if analysis.completed_at else None,
        "error_message": analysis.error_message,
    }

    # Add results if completed
    if analysis.status == "completed":
        status_data.update(
            {
                "duration": analysis.duration,
                "total_speakers": analysis.total_speakers,
                "total_segments": analysis.total_segments,
                "total_discrepancies": analysis.total_discrepancies,
                "critical_discrepancies": analysis.critical_discrepancies,
            }
        )

    return jsonify(status_data)


@app.route("/api/analysis/<analysis_id>/report/<format>", methods=["GET"])
@login_required
def download_report(analysis_id, format):
    """Download analysis report"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    if analysis.status != "completed":
        return jsonify({"error": "Analysis not completed"}), 400

    # Get report file path
    if format == "json":
        report_path = analysis.report_json_path
        mimetype = "application/json"
    elif format == "txt":
        report_path = analysis.report_txt_path
        mimetype = "text/plain"
    elif format == "md":
        report_path = analysis.report_md_path
        mimetype = "text/markdown"
    elif format in ["pdf", "docx"]:
        # Generate on-demand for PDF/DOCX
        return export_analysis_report(analysis, format)
    else:
        return jsonify({"error": "Invalid format"}), 400

    if not report_path or not Path(report_path).exists():
        return jsonify({"error": "Report file not found"}), 404

    # Log audit
    AuditLog.log("report_downloaded", "analysis", analysis_id, {"format": format})

    return send_file(
        report_path,
        mimetype=mimetype,
        as_attachment=True,
        download_name=f"BWC_Analysis_{analysis.case_number or analysis_id}.{format}",
    )


def export_analysis_report(analysis, format):
    """Export analysis to PDF or DOCX format"""
    try:
        # Load the JSON report
        if not analysis.report_json_path or not Path(analysis.report_json_path).exists():
            return jsonify({"error": "Analysis report not available"}), 404

        import json

        with open(analysis.report_json_path, "r") as f:
            json.load(f)

        output_dir = Path(app.config["ANALYSIS_FOLDER"]) / analysis.id
        output_dir.mkdir(parents=True, exist_ok=True)

        if format == "pdf":
            # Generate PDF report
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import (Paragraph, SimpleDocTemplate,
                                            Spacer, Table, TableStyle)

            pdf_path = output_dir / f"report_{analysis.id}.pdf"
            doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Title
            story.append(Paragraph("<b>BWC Forensic Analysis Report</b>", styles["Title"]))
            story.append(Spacer(1, 12))

            # Case Information
            story.append(Paragraph("<b>Case Information</b>", styles["Heading2"]))
            case_info = [
                ["Case Number:", analysis.case_number or "N/A"],
                ["Evidence Number:", analysis.evidence_number or "N/A"],
                ["Filename:", analysis.filename],
                ["File Hash (SHA-256):", analysis.file_hash],
                ["Analysis Date:", analysis.created_at.strftime("%Y-%m-%d %H:%M:%S")],
                ["Duration:", f"{analysis.duration:.2f} seconds" if analysis.duration else "N/A"],
            ]
            case_table = Table(case_info)
            case_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                        ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )
            story.append(case_table)
            story.append(Spacer(1, 20))

            # Results Summary
            story.append(Paragraph("<b>Analysis Results</b>", styles["Heading2"]))
            results = [
                ["Total Speakers:", str(analysis.total_speakers or 0)],
                ["Transcript Segments:", str(analysis.total_segments or 0)],
                ["Total Discrepancies:", str(analysis.total_discrepancies or 0)],
                ["Critical Discrepancies:", str(analysis.critical_discrepancies or 0)],
            ]
            results_table = Table(results)
            results_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ]
                )
            )
            story.append(results_table)
            story.append(Spacer(1, 20))

            # Chain of Custody
            story.append(Paragraph("<b>Chain of Custody</b>", styles["Heading2"]))
            story.append(
                Paragraph(f"File Integrity (SHA-256): {analysis.file_hash}", styles["Normal"])
            )
            story.append(
                Paragraph(
                    f"Acquired By: {analysis.acquired_by or 'Not specified'}", styles["Normal"]
                )
            )
            story.append(
                Paragraph(f"Source: {analysis.source or 'Not specified'}", styles["Normal"])
            )

            doc.build(story)

            # Log audit
            AuditLog.log("report_exported", "analysis", analysis.id, {"format": "pdf"})

            return send_file(
                str(pdf_path),
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"BWC_Analysis_{analysis.case_number or analysis.id}.pdf",
            )

        elif format == "docx":
            # Generate DOCX report
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            docx_path = output_dir / f"report_{analysis.id}.docx"
            doc = Document()

            # Title
            title = doc.add_heading("BWC Forensic Analysis Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Case Information
            doc.add_heading("Case Information", level=1)
            table = doc.add_table(rows=6, cols=2)
            table.style = "Light Grid Accent 1"

            cells = table.rows[0].cells
            cells[0].text = "Case Number:"
            cells[1].text = analysis.case_number or "N/A"

            cells = table.rows[1].cells
            cells[0].text = "Evidence Number:"
            cells[1].text = analysis.evidence_number or "N/A"

            cells = table.rows[2].cells
            cells[0].text = "Filename:"
            cells[1].text = analysis.filename

            cells = table.rows[3].cells
            cells[0].text = "File Hash (SHA-256):"
            cells[1].text = analysis.file_hash

            cells = table.rows[4].cells
            cells[0].text = "Analysis Date:"
            cells[1].text = analysis.created_at.strftime("%Y-%m-%d %H:%M:%S")

            cells = table.rows[5].cells
            cells[0].text = "Duration:"
            cells[1].text = f"{analysis.duration:.2f} seconds" if analysis.duration else "N/A"

            # Results Summary
            doc.add_heading("Analysis Results", level=1)
            results_table = doc.add_table(rows=4, cols=2)
            results_table.style = "Light Grid Accent 1"

            cells = results_table.rows[0].cells
            cells[0].text = "Total Speakers:"
            cells[1].text = str(analysis.total_speakers or 0)

            cells = results_table.rows[1].cells
            cells[0].text = "Transcript Segments:"
            cells[1].text = str(analysis.total_segments or 0)

            cells = results_table.rows[2].cells
            cells[0].text = "Total Discrepancies:"
            cells[1].text = str(analysis.total_discrepancies or 0)

            cells = results_table.rows[3].cells
            cells[0].text = "Critical Discrepancies:"
            cells[1].text = str(analysis.critical_discrepancies or 0)

            # Chain of Custody
            doc.add_heading("Chain of Custody", level=1)
            doc.add_paragraph(f"File Integrity (SHA-256): {analysis.file_hash}")
            doc.add_paragraph(f'Acquired By: {analysis.acquired_by or "Not specified"}')
            doc.add_paragraph(f'Source: {analysis.source or "Not specified"}')

            # Attribution
            doc.add_heading("Attribution", level=1)
            doc.add_paragraph("This analysis was generated using:")
            doc.add_paragraph("• OpenAI Whisper (Apache 2.0 License) - Audio transcription")
            doc.add_paragraph("• pyannote.audio (MIT License) - Speaker diarization")
            doc.add_paragraph("• BarberX Legal Technologies - Forensic analysis platform")

            doc.save(str(docx_path))

            # Log audit
            AuditLog.log("report_exported", "analysis", analysis.id, {"format": "docx"})

            return send_file(
                str(docx_path),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=f"BWC_Analysis_{analysis.case_number or analysis.id}.docx",
            )

        elif format == "json":
            # Generate JSON export
            import json

            json_path = output_dir / f"report_{analysis.id}.json"

            json_data = {
                "case_information": {
                    "case_number": analysis.case_number or "N/A",
                    "evidence_number": analysis.evidence_number or "N/A",
                    "filename": analysis.filename,
                    "file_hash": analysis.file_hash,
                    "file_size": analysis.file_size,
                    "duration": analysis.duration,
                    "analysis_date": analysis.created_at.isoformat(),
                    "updated_at": analysis.updated_at.isoformat(),
                },
                "analysis_results": {
                    "status": analysis.status,
                    "progress": analysis.progress,
                    "current_step": analysis.current_step,
                    "total_speakers": analysis.total_speakers or 0,
                    "total_segments": analysis.total_segments or 0,
                    "total_discrepancies": analysis.total_discrepancies or 0,
                    "critical_discrepancies": analysis.critical_discrepancies or 0,
                },
                "chain_of_custody": {
                    "file_integrity_hash": analysis.file_hash,
                    "acquired_by": analysis.acquired_by or "Not specified",
                    "source": analysis.source or "Not specified",
                    "analyzed_by": current_user.email,
                },
                "metadata": analysis.metadata or {},
                "export_timestamp": datetime.utcnow().isoformat(),
                "platform": "BarberX Legal Tech Platform",
                "version": "2.0",
            }

            with open(str(json_path), "w") as f:
                json.dump(json_data, f, indent=2)

            # Log audit
            AuditLog.log("report_exported", "analysis", analysis.id, {"format": "json"})

            return send_file(
                str(json_path),
                mimetype="application/json",
                as_attachment=True,
                download_name=f"BWC_Analysis_{analysis.case_number or analysis.id}.json",
            )

        elif format == "txt":
            # Generate plain text export
            txt_path = output_dir / f"report_{analysis.id}.txt"

            duration_str = (
                f"{int(analysis.duration // 60)}m {int(analysis.duration % 60)}s"
                if analysis.duration
                else "N/A"
            )
            file_size_str = (
                f"{analysis.file_size / (1024*1024):.2f} MB" if analysis.file_size else "N/A"
            )

            text_content = f"""
========================================
BWC FORENSIC ANALYSIS REPORT
========================================
Generated: {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')}

CASE INFORMATION
----------------
Case Number: {analysis.case_number or 'N/A'}
Evidence Number: {analysis.evidence_number or 'N/A'}
Filename: {analysis.filename}
File Size: {file_size_str}
Duration: {duration_str}
SHA-256 Hash: {analysis.file_hash}

ANALYSIS RESULTS
----------------
Speakers Identified: {analysis.total_speakers or 0}
Transcript Segments: {analysis.total_segments or 0}
Total Discrepancies: {analysis.total_discrepancies or 0}
Critical Issues: {analysis.critical_discrepancies or 0}
Analysis Status: {analysis.status.upper()}
Completion Progress: {analysis.progress}%
Current Step: {analysis.current_step or 'N/A'}

CHAIN OF CUSTODY
----------------
Evidence Acquired By: {analysis.acquired_by or 'Not specified'}
Source/Origin: {analysis.source or 'Not specified'}
Upload Date: {analysis.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')}
Analysis Completed: {analysis.updated_at.strftime('%Y-%m-%d %H:%M:%S UTC')}
Analyzed By: {current_user.email}

INTEGRITY VERIFICATION
----------------------
File Hash (SHA-256): {analysis.file_hash}
This cryptographic hash ensures file integrity and admissibility
in legal proceedings. Any modification to the original file will
result in a different hash value, allowing detection of tampering.

LEGAL NOTICE
------------
This report is generated for official use only and contains chain-of-custody
verified evidence. The SHA-256 cryptographic hash ensures file integrity and
admissibility. Any unauthorized modification, reproduction, or distribution
is prohibited.

This analysis was conducted using BarberX Legal Tech Platform certified
forensic analysis tools in compliance with digital evidence standards.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BarberX Legal Tech Platform
BWC Forensic Analysis System | Copyright © 2024-2026
For official use only - Confidential
"""

            with open(str(txt_path), "w") as f:
                f.write(text_content)

            # Log audit
            AuditLog.log("report_exported", "analysis", analysis.id, {"format": "txt"})

            return send_file(
                str(txt_path),
                mimetype="text/plain",
                as_attachment=True,
                download_name=f"BWC_Analysis_{analysis.case_number or analysis.id}.txt",
            )

        elif format == "md":
            # Generate Markdown export
            md_path = output_dir / f"report_{analysis.id}.md"

            duration_str = (
                f"{int(analysis.duration // 60)}m {int(analysis.duration % 60)}s"
                if analysis.duration
                else "N/A"
            )
            file_size_str = (
                f"{analysis.file_size / (1024*1024):.2f} MB" if analysis.file_size else "N/A"
            )

            markdown_content = f"""# BWC FORENSIC ANALYSIS REPORT

**Generated:** {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')}

---

## EXECUTIVE SUMMARY

This report presents the forensic analysis results for body-worn camera footage case **{analysis.case_number or 'N/A'}**.
The analysis identified **{analysis.total_speakers or 0}** distinct speaker(s) across **{analysis.total_segments or 0}** transcript segments.

{"**⚠️ ALERT:** " + str(analysis.critical_discrepancies) + " critical discrepancies require immediate attention." if analysis.critical_discrepancies and analysis.critical_discrepancies > 0 else "✅ No critical issues were detected."}

---

## CASE INFORMATION

| Field | Value |
|-------|-------|
| Case Number | {analysis.case_number or 'N/A'} |
| Evidence Number | {analysis.evidence_number or 'N/A'} |
| Filename | {analysis.filename} |
| File Size | {file_size_str} |
| Duration | {duration_str} |
| SHA-256 Hash | `{analysis.file_hash}` |

---

## ANALYSIS RESULTS

| Metric | Value | Status |
|--------|-------|--------|
| Speakers Identified | {analysis.total_speakers or 0} | ✓ |
| Transcript Segments | {analysis.total_segments or 0} | ✓ |
| Total Discrepancies | {analysis.total_discrepancies or 0} | {'⚠️' if analysis.total_discrepancies and analysis.total_discrepancies > 0 else '✓'} |
| Critical Issues | {analysis.critical_discrepancies or 0} | {'⚠️' if analysis.critical_discrepancies and analysis.critical_discrepancies > 0 else '✓'} |
| Analysis Status | {analysis.status.upper()} | {'✓' if analysis.status == 'completed' else '⏳'} |
| Completion | {analysis.progress}% | {'✓' if analysis.progress == 100 else '⏳'} |

---

## CHAIN OF CUSTODY

| Event | Details |
|-------|---------|
| Evidence Acquired By | {analysis.acquired_by or 'Not specified'} |
| Source/Origin | {analysis.source or 'Not specified'} |
| Upload Date | {analysis.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')} |
| Analysis Completed | {analysis.updated_at.strftime('%Y-%m-%d %H:%M:%S UTC')} |
| Analyzed By | {current_user.email} |
| Integrity Verification | `SHA-256: {analysis.file_hash}` |

---

## INTEGRITY VERIFICATION

```
File Hash (SHA-256): {analysis.file_hash}
```

This cryptographic hash ensures file integrity and admissibility in legal proceedings.
Any modification to the original file will result in a different hash value, allowing
detection of tampering.

---

## LEGAL NOTICE

This report is generated for official use only and contains chain-of-custody verified evidence.
The **SHA-256 cryptographic hash** ensures file integrity and admissibility.

⚠️ Any unauthorized modification, reproduction, or distribution is prohibited.

This analysis was conducted using BarberX Legal Tech Platform certified forensic analysis tools
in compliance with digital evidence standards.

---

**BarberX Legal Tech Platform**
BWC Forensic Analysis System | Copyright © 2024-2026
*For official use only - Confidential*
"""

            with open(str(md_path), "w") as f:
                f.write(markdown_content)

            # Log audit
            AuditLog.log("report_exported", "analysis", analysis.id, {"format": "md"})

            return send_file(
                str(md_path),
                mimetype="text/markdown",
                as_attachment=True,
                download_name=f"BWC_Analysis_{analysis.case_number or analysis.id}.md",
            )

    except ImportError:
        logger.error("Export dependencies not installed", exc_info=True)
        return (
            jsonify(
                {
                    "error": "Export dependencies not installed",
                    "message": "Please install: pip install reportlab python-docx",
                }
            ),
            500,
        )
    except Exception as e:
        app.logger.error(f"Export error: {str(e)}")
        return jsonify({"error": f"Export failed: {str(e)}"}), 500


# ========================================
# API ROUTES - User Management
# ========================================


@app.route("/api/user/profile", methods=["GET", "PUT"])
@login_required
def user_profile():
    """Get or update user profile"""
    if request.method == "GET":
        return jsonify(current_user.to_dict())

    # PUT - update profile
    data = request.get_json()

    if "full_name" in data:
        current_user.full_name = data["full_name"]
    if "organization" in data:
        current_user.organization = data["organization"]

    db.session.commit()

    # Log audit
    AuditLog.log("profile_updated", "user", str(current_user.id))

    return jsonify(current_user.to_dict())


@app.route("/api/user/api-keys", methods=["GET", "POST"])
@login_required
def manage_api_keys():
    """Manage API keys"""
    if request.method == "GET":
        keys = current_user.api_keys.filter_by(is_active=True).all()
        return jsonify({"api_keys": [k.to_dict() for k in keys]})

    # Create new API key
    data = request.get_json()

    if not data.get("name"):
        return jsonify({"error": "API key name required"}), 400

    # Check tier limits
    if current_user.subscription_tier == "free":
        return jsonify({"error": "API access requires Professional or Enterprise plan"}), 403

    api_key = APIKey(user_id=current_user.id, name=data["name"])
    api_key.generate_key()

    db.session.add(api_key)
    db.session.commit()

    # Log audit
    AuditLog.log("api_key_created", "api_key", str(api_key.id))

    return jsonify(api_key.to_dict()), 201


# ========================================
# ADMIN ROUTES
# ========================================

# ========================================
# DASHBOARD API ROUTES
# ========================================


@app.route("/api/dashboard-stats", methods=["GET"])
@login_required
def dashboard_stats():
    """Get dashboard statistics for current user"""
    from datetime import timedelta

    from sqlalchemy import func

    # Get current month's analysis count
    now = datetime.utcnow()
    first_day = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    analyses_this_month = Analysis.query.filter(
        Analysis.user_id == current_user.id, Analysis.created_at >= first_day
    ).count()

    # Get tier limits
    tier_limits = current_user.get_tier_limits()

    # Get completed count
    completed_count = Analysis.query.filter_by(user_id=current_user.id, status="completed").count()

    # Get analyzing count
    analyzing_count = Analysis.query.filter_by(user_id=current_user.id, status="analyzing").count()

    # Get failed count
    failed_count = Analysis.query.filter_by(user_id=current_user.id, status="failed").count()

    # Get daily activity for last 7 days
    seven_days_ago = now - timedelta(days=7)
    daily_activity = (
        db.session.query(
            func.date(Analysis.created_at).label("date"), func.count(Analysis.id).label("count")
        )
        .filter(Analysis.user_id == current_user.id, Analysis.created_at >= seven_days_ago)
        .group_by(func.date(Analysis.created_at))
        .all()
    )

    # Format daily activity
    daily_counts = [0] * 7
    for activity in daily_activity:
        days_diff = (now.date() - activity.date).days
        if 0 <= days_diff < 7:
            daily_counts[6 - days_diff] = activity.count

    return jsonify(
        {
            "analyses_this_month": analyses_this_month,
            "storage_used_mb": current_user.storage_used_mb,
            "tier_limits": tier_limits,
            "completed_count": completed_count,
            "analyzing_count": analyzing_count,
            "failed_count": failed_count,
            "daily_activity": daily_counts,
            "subscription_tier": current_user.tier.name.lower(),
            "tier_name": current_user.tier_name,
            "is_admin": current_user.is_admin,
        }
    )


@app.route("/api/analysis/<analysis_id>", methods=["GET"])
@login_required
def get_analysis_details(analysis_id):
    """Get specific analysis details"""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()

    if not analysis:
        return jsonify({"error": "Analysis not found"}), 404

    return jsonify(analysis.to_dict())


@app.route("/api/subscription/upgrade", methods=["POST"])
@login_required
def upgrade_subscription():
    """Upgrade user subscription tier"""
    data = request.get_json()
    new_tier = data.get("tier")

    valid_tiers = ["free", "professional", "enterprise"]

    if new_tier not in valid_tiers:
        return jsonify({"error": "Invalid tier"}), 400

    # In production, integrate with Stripe here
    current_user.subscription_tier = new_tier
    db.session.commit()

    # Log audit
    AuditLog.log(
        "subscription_upgraded",
        "user",
        str(current_user.id),
        {"from": current_user.subscription_tier, "to": new_tier},
    )

    return jsonify(
        {
            "message": "Subscription upgraded successfully",
            "new_tier": new_tier,
            "tier_limits": current_user.get_tier_limits(),
        }
    )


@app.route("/api/user/api-keys", methods=["GET"])
@login_required
def list_api_keys():
    """List user's API keys"""
    if current_user.subscription_tier not in ["professional", "enterprise"]:
        return jsonify({"error": "API keys require Professional or Enterprise tier"}), 403

    keys = APIKey.query.filter_by(user_id=current_user.id).all()

    return jsonify({"total": len(keys), "keys": [k.to_dict() for k in keys]})


@app.route("/api/user/api-keys/<key_id>", methods=["DELETE"])
@login_required
def delete_api_key(key_id):
    """Delete an API key"""
    key = APIKey.query.filter_by(id=key_id, user_id=current_user.id).first()

    if not key:
        return jsonify({"error": "API key not found"}), 404

    db.session.delete(key)
    db.session.commit()

    # Log audit
    AuditLog.log("api_key_deleted", "api_key", str(key_id))

    return jsonify({"message": "API key deleted successfully"})


@app.route("/api/audit-logs", methods=["GET"])
@login_required
def list_audit_logs():
    """List user's audit logs"""
    limit = request.args.get("limit", 50, type=int)

    logs = (
        AuditLog.query.filter_by(user_id=current_user.id)
        .order_by(AuditLog.created_at.desc())
        .limit(limit)
        .all()
    )

    return jsonify(
        {
            "total": len(logs),
            "logs": [
                {
                    "id": log.id,
                    "action": log.action,
                    "resource_type": log.resource_type,
                    "resource_id": log.resource_id,
                    "details": log.details,
                    "ip_address": log.ip_address,
                    "created_at": log.created_at.isoformat(),
                }
                for log in logs
            ],
        }
    )


# ========================================
# ADMIN ROUTES
# ========================================


@app.route("/admin/users", methods=["GET"])
@login_required
def admin_list_users():
    """Admin: List all users with pagination"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    # Add pagination
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 50, type=int)
    per_page = min(per_page, 100)  # Cap at 100

    users_query = User.query.order_by(User.created_at.desc())
    total = users_query.count()
    users = users_query.limit(per_page).offset((page - 1) * per_page).all()

    return jsonify(
        {"total": total, "page": page, "per_page": per_page, "users": [u.to_dict() for u in users]}
    )


@app.route("/admin/users/<int:user_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def admin_manage_user(user_id):
    """Admin: Get, update, or delete specific user"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    user = User.query.get_or_404(user_id)

    if request.method == "GET":
        return jsonify(user.to_dict())

    elif request.method == "PUT":
        data = request.get_json()

        # Update allowed fields
        if "full_name" in data:
            user.full_name = data["full_name"]
        if "organization" in data:
            user.organization = data["organization"]
        if "subscription_tier" in data:
            if data["subscription_tier"] in ["free", "professional", "enterprise"]:
                user.subscription_tier = data["subscription_tier"]
        if "role" in data:
            if data["role"] in ["user", "pro", "admin"]:
                user.role = data["role"]
        if "is_active" in data:
            user.is_active = data["is_active"]
        if "is_verified" in data:
            user.is_verified = data["is_verified"]

        db.session.commit()

        # Log audit
        AuditLog.log("admin_user_updated", "user", str(user_id), {"updated_by": current_user.email})

        return jsonify({"message": "User updated successfully", "user": user.to_dict()})

    elif request.method == "DELETE":
        # Don't allow deleting yourself
        if user.id == current_user.id:
            return jsonify({"error": "Cannot delete your own account"}), 400

        db.session.delete(user)
        db.session.commit()

        # Log audit
        AuditLog.log(
            "admin_user_deleted",
            "user",
            str(user_id),
            {"deleted_by": current_user.email, "deleted_email": user.email},
        )

        return jsonify({"message": "User deleted successfully"})


@app.route("/admin/users/<int:user_id>/toggle-status", methods=["POST"])
@login_required
def admin_toggle_user_status(user_id):
    """Admin: Enable/disable user account"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    user = User.query.get_or_404(user_id)

    # Don't allow disabling yourself
    if user.id == current_user.id:
        return jsonify({"error": "Cannot disable your own account"}), 400

    user.is_active = not user.is_active
    db.session.commit()

    # Log audit
    AuditLog.log(
        "admin_user_status_toggled",
        "user",
        str(user_id),
        {
            "new_status": "active" if user.is_active else "inactive",
            "changed_by": current_user.email,
        },
    )

    return jsonify(
        {
            "message": f'User {"activated" if user.is_active else "deactivated"} successfully',
            "is_active": user.is_active,
        }
    )


@app.route("/admin/users/<int:user_id>/reset-password", methods=["POST"])
@login_required
def admin_reset_user_password(user_id):
    """Admin: Reset user password"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    user = User.query.get_or_404(user_id)
    data = request.get_json()

    if not data.get("new_password"):
        return jsonify({"error": "New password required"}), 400

    user.set_password(data["new_password"])
    db.session.commit()

    # Log audit
    AuditLog.log("admin_password_reset", "user", str(user_id), {"reset_by": current_user.email})

    return jsonify({"message": "Password reset successfully"})


@app.route("/admin/analyses", methods=["GET"])
@login_required
def admin_list_analyses():
    """Admin: List all analyses"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    status_filter = request.args.get("status")
    limit = request.args.get("limit", 50, type=int)
    offset = request.args.get("offset", 0, type=int)

    query = Analysis.query

    if status_filter:
        query = query.filter_by(status=status_filter)

    total = query.count()
    analyses = query.order_by(Analysis.created_at.desc()).limit(limit).offset(offset).all()

    return jsonify(
        {
            "total": total,
            "limit": limit,
            "offset": offset,
            "analyses": [a.to_dict() for a in analyses],
        }
    )


@app.route("/admin/analyses/<analysis_id>", methods=["DELETE"])
@login_required
def admin_delete_analysis(analysis_id):
    """Admin: Delete an analysis"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    analysis = Analysis.query.get_or_404(analysis_id)

    # Delete associated file if exists
    if analysis.file_path and os.path.exists(analysis.file_path):
        try:
            os.remove(analysis.file_path)
        except Exception as e:
            app.logger.error(f"Error deleting file: {e}")

    # Update user storage
    if analysis.user:
        analysis.user.storage_used_mb -= analysis.file_size / (1024 * 1024)
        if analysis.user.storage_used_mb < 0:
            analysis.user.storage_used_mb = 0

    db.session.delete(analysis)
    db.session.commit()

    # Log audit
    AuditLog.log(
        "admin_analysis_deleted", "analysis", analysis_id, {"deleted_by": current_user.email}
    )

    return jsonify({"message": "Analysis deleted successfully"})


@app.route("/admin/stats", methods=["GET"])
@login_required
def admin_stats():
    """Admin: Get platform statistics (optimized with aggregation queries)"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    from datetime import timedelta

    from sqlalchemy import func

    # Use single queries with aggregation instead of multiple count queries
    total_users = User.query.count()
    total_analyses = Analysis.query.count()

    # Get status counts in a single query
    status_counts = dict(
        db.session.query(Analysis.status, func.count(Analysis.id)).group_by(Analysis.status).all()
    )
    completed_analyses = status_counts.get("completed", 0)
    analyzing = status_counts.get("analyzing", 0)
    failed = status_counts.get("failed", 0)

    # Get subscription tier counts in a single query

    tier_counts = {}
    if hasattr(User, "tier"):
        tier_result = db.session.query(User.tier, func.count(User.id)).group_by(User.tier).all()
        tier_counts = {
            tier.name if hasattr(tier, "name") else str(tier): count for tier, count in tier_result
        }

    free_users = tier_counts.get("FREE", 0)
    pro_users = tier_counts.get("PROFESSIONAL", 0)
    premium_users = tier_counts.get("PREMIUM", 0)
    enterprise_users = tier_counts.get("ENTERPRISE", 0)

    # Storage stats
    total_storage = db.session.query(func.sum(User.storage_used_mb)).scalar() or 0

    # Active users (logged in last 30 days)
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    active_users = User.query.filter(User.last_login >= thirty_days_ago).count()

    # Daily activity for last 7 days (optimized with group by)
    seven_days_ago = datetime.utcnow() - timedelta(days=7)
    daily_analyses = (
        db.session.query(
            func.date(Analysis.created_at).label("date"), func.count(Analysis.id).label("count")
        )
        .filter(Analysis.created_at >= seven_days_ago)
        .group_by(func.date(Analysis.created_at))
        .all()
    )

    # Format daily activity
    now = datetime.utcnow()
    daily_counts = [0] * 7
    for activity in daily_analyses:
        days_diff = (now.date() - activity.date).days
        if 0 <= days_diff < 7:
            daily_counts[6 - days_diff] = activity.count

    # Estimated revenue (based on tier pricing)
    tier_pricing = {"FREE": 0, "PROFESSIONAL": 49, "PREMIUM": 99, "ENTERPRISE": 249}
    estimated_revenue = sum(
        tier_pricing.get(tier, 0) * count for tier, count in tier_counts.items()
    )

    return jsonify(
        {
            "total_users": total_users,
            "active_users": active_users,
            "total_analyses": total_analyses,
            "total_evidence": total_analyses,  # Alias for evidence count
            "completed_analyses": completed_analyses,
            "analyzing_count": analyzing,
            "failed_count": failed,
            "success_rate": (
                (completed_analyses / total_analyses * 100) if total_analyses > 0 else 0
            ),
            "subscription_breakdown": {
                "free": free_users,
                "professional": pro_users,
                "premium": premium_users,
                "enterprise": enterprise_users,
            },
            "tier_distribution": {
                "FREE": free_users,
                "PROFESSIONAL": pro_users,
                "PREMIUM": premium_users,
                "ENTERPRISE": enterprise_users,
            },
            "revenue": estimated_revenue,
            "total_storage_gb": round(total_storage / 1024, 2),
            "daily_activity": daily_counts,
            "total_agents": 0,  # Placeholder for agent stats
            "active_agents": 0,
            "completed_jobs_today": completed_analyses,
        }
    )


@app.route("/admin/operations-summary", methods=["GET"])
@login_required
def admin_operations_summary():
    """Admin: Operational costs, risk flags, and feature toggles"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    from datetime import timedelta

    import psutil
    from sqlalchemy import func

    def ensure_setting(key, default, value_type="string", category="operations", description=""):
        setting = AppSettings.query.filter_by(key=key).first()
        if not setting:
            setting = AppSettings(
                key=key,
                value=str(default),
                value_type=value_type,
                category=category,
                description=description,
                updated_by=current_user.id,
            )
            db.session.add(setting)
            db.session.commit()
        return AppSettings.get(key, default)

    # Cost settings
    cost_per_gb = ensure_setting(
        "ops.cost_per_gb",
        0.023,
        "float",
        description="Estimated monthly storage cost per GB",
    )
    cost_per_analysis = ensure_setting(
        "ops.cost_per_analysis",
        0.05,
        "float",
        description="Estimated processing cost per analysis",
    )
    cost_per_active_user = ensure_setting(
        "ops.cost_per_active_user",
        0.10,
        "float",
        description="Estimated monthly cost per active user",
    )
    monthly_budget = ensure_setting(
        "ops.monthly_budget",
        250,
        "float",
        description="Monthly operations budget",
    )

    # Thresholds
    disk_warn = ensure_setting("ops.disk_warn_pct", 85, "int")
    disk_crit = ensure_setting("ops.disk_crit_pct", 95, "int")
    memory_warn = ensure_setting("ops.memory_warn_pct", 85, "int")
    memory_crit = ensure_setting("ops.memory_crit_pct", 95, "int")
    cpu_warn = ensure_setting("ops.cpu_warn_pct", 85, "int")
    cpu_crit = ensure_setting("ops.cpu_crit_pct", 95, "int")
    failed_warn = ensure_setting("ops.failed_warn_count", 5, "int")
    failed_crit = ensure_setting("ops.failed_crit_count", 20, "int")
    backlog_warn = ensure_setting("ops.backlog_warn_count", 10, "int")
    backlog_crit = ensure_setting("ops.backlog_crit_count", 30, "int")
    storage_warn_gb = ensure_setting("ops.storage_warn_gb", 50, "int")
    storage_crit_gb = ensure_setting("ops.storage_crit_gb", 200, "int")
    budget_warn = ensure_setting("ops.budget_warn_pct", 80, "int")
    budget_crit = ensure_setting("ops.budget_crit_pct", 100, "int")

    # Usage data
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    seven_days_ago = datetime.utcnow() - timedelta(days=7)

    analyses_last_30 = Analysis.query.filter(Analysis.created_at >= thirty_days_ago).count()
    failed_last_7 = Analysis.query.filter(
        Analysis.created_at >= seven_days_ago, Analysis.status == "failed"
    ).count()
    analyzing_now = Analysis.query.filter(Analysis.status == "analyzing").count()

    active_users = User.query.filter(User.last_login >= thirty_days_ago).count()

    # Storage sizes
    instance_path = app.instance_path
    db_paths = [
        os.path.join(instance_path, "barberx_FRESH.db"),
        os.path.join(instance_path, "barberx_legal.db"),
    ]

    db_size_mb = sum(
        os.path.getsize(path) / (1024 * 1024) for path in db_paths if os.path.exists(path)
    )

    upload_folder = app.config.get("UPLOAD_FOLDER", "")
    upload_size_mb = 0
    if upload_folder and os.path.exists(upload_folder):
        for root, dirs, files in os.walk(upload_folder):
            upload_size_mb += sum(os.path.getsize(os.path.join(root, f)) for f in files)
        upload_size_mb = upload_size_mb / (1024 * 1024)

    total_storage_gb = (db_size_mb + upload_size_mb) / 1024

    # System metrics
    cpu_percent = psutil.cpu_percent(interval=0.4)
    memory = psutil.virtual_memory()
    disk = psutil.disk_usage("/")

    # Cost calculations
    storage_cost = total_storage_gb * cost_per_gb
    processing_cost = analyses_last_30 * cost_per_analysis
    active_users_cost = active_users * cost_per_active_user
    total_monthly = storage_cost + processing_cost + active_users_cost
    budget_utilization = (total_monthly / monthly_budget * 100) if monthly_budget else 0

    # Risk flags
    risk_flags = []

    def add_flag(level, label, message):
        risk_flags.append({"level": level, "label": label, "message": message})

    if disk.percent >= disk_crit:
        add_flag("critical", "Disk Usage", f"Disk at {disk.percent:.1f}% (critical)")
    elif disk.percent >= disk_warn:
        add_flag("warning", "Disk Usage", f"Disk at {disk.percent:.1f}% (warning)")

    if memory.percent >= memory_crit:
        add_flag("critical", "Memory Usage", f"Memory at {memory.percent:.1f}% (critical)")
    elif memory.percent >= memory_warn:
        add_flag("warning", "Memory Usage", f"Memory at {memory.percent:.1f}% (warning)")

    if cpu_percent >= cpu_crit:
        add_flag("critical", "CPU Load", f"CPU at {cpu_percent:.1f}% (critical)")
    elif cpu_percent >= cpu_warn:
        add_flag("warning", "CPU Load", f"CPU at {cpu_percent:.1f}% (warning)")

    if failed_last_7 >= failed_crit:
        add_flag("critical", "Failed Analyses", f"{failed_last_7} failures in 7 days")
    elif failed_last_7 >= failed_warn:
        add_flag("warning", "Failed Analyses", f"{failed_last_7} failures in 7 days")

    if analyzing_now >= backlog_crit:
        add_flag("critical", "Processing Backlog", f"{analyzing_now} analyses in queue")
    elif analyzing_now >= backlog_warn:
        add_flag("warning", "Processing Backlog", f"{analyzing_now} analyses in queue")

    if total_storage_gb >= storage_crit_gb:
        add_flag("critical", "Storage Growth", f"{total_storage_gb:.1f} GB used")
    elif total_storage_gb >= storage_warn_gb:
        add_flag("warning", "Storage Growth", f"{total_storage_gb:.1f} GB used")

    if budget_utilization >= budget_crit:
        add_flag("critical", "Budget Risk", f"{budget_utilization:.1f}% of monthly budget used")
    elif budget_utilization >= budget_warn:
        add_flag("warning", "Budget Risk", f"{budget_utilization:.1f}% of monthly budget used")

    # Feature toggles
    toggle_defs = [
        ("ops.toggle.ai_chat", True, "AI Chat Assistant"),
        ("ops.toggle.video_processing", True, "Video Processing"),
        ("ops.toggle.document_ocr", True, "Document OCR"),
        ("ops.toggle.batch_uploads", True, "Batch Uploads"),
        ("ops.toggle.email_notifications", True, "Email Notifications"),
    ]
    toggles = []
    for key, default, label in toggle_defs:
        value = ensure_setting(key, default, "bool", description=label)
        toggles.append({"key": key, "label": label, "value": bool(value)})

    return jsonify(
        {
            "costs": {
                "total_monthly": round(total_monthly, 2),
                "storage_monthly": round(storage_cost, 2),
                "processing_monthly": round(processing_cost, 2),
                "active_users_monthly": round(active_users_cost, 2),
                "monthly_budget": round(monthly_budget, 2),
                "budget_utilization_pct": round(budget_utilization, 2),
            },
            "usage": {
                "storage_gb": round(total_storage_gb, 2),
                "analyses_last_30_days": analyses_last_30,
                "failed_last_7_days": failed_last_7,
                "analyzing_now": analyzing_now,
                "active_users": active_users,
                "cpu_percent": round(cpu_percent, 2),
                "memory_percent": round(memory.percent, 2),
                "disk_percent": round(disk.percent, 2),
            },
            "risk_flags": risk_flags,
            "toggles": toggles,
        }
    )


@app.route("/admin/operations-toggles", methods=["POST"])
@login_required
def admin_operations_toggles():
    """Admin: Update operational feature toggles"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    data = request.get_json() or {}
    key = data.get("key")
    value = data.get("value")

    if not key or not key.startswith("ops.toggle."):
        return jsonify({"error": "Invalid toggle key"}), 400

    AppSettings.set(
        key,
        bool(value),
        value_type="bool",
        category="operations",
        description="Operations toggle",
    )

    AuditLog.log("operations_toggle", "AppSettings", key, {"value": bool(value)})

    return jsonify({"message": "Toggle updated", "key": key, "value": bool(value)})


@app.route("/admin/site-controls", methods=["GET", "POST"])
@login_required
def admin_site_controls():
    """Admin: Manage site-wide UI and maintenance controls"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    allowed_keys = {
        "ui.theme": ("system", "string", "Site theme: system, light, dark"),
        "ui.brand_color": ("#c41e3a", "string", "Primary brand color"),
        "ui.accent_color": ("#1e40af", "string", "Accent color"),
        "site.maintenance_mode": (False, "bool", "Enable maintenance mode"),
        "site.maintenance_message": (
            "We are performing maintenance. Please check back shortly.",
            "string",
            "Maintenance banner message",
        ),
        "site.banner_enabled": (False, "bool", "Enable announcement banner"),
        "site.banner_message": ("", "string", "Announcement banner message"),
        "site.allow_signup": (True, "bool", "Allow new user signups"),
        "site.footer_notice": ("", "string", "Footer notice text"),
        "site.chat_enabled": (True, "bool", "Enable chat interface"),
    }

    def ensure_setting(key, default, value_type, description):
        setting = AppSettings.query.filter_by(key=key).first()
        if not setting:
            setting = AppSettings(
                key=key,
                value=str(default),
                value_type=value_type,
                category="ui",
                description=description,
                updated_by=current_user.id,
            )
            db.session.add(setting)
            db.session.commit()
        return AppSettings.get(key, default)

    if request.method == "GET":
        values = {
            key: ensure_setting(key, default, value_type, description)
            for key, (default, value_type, description) in allowed_keys.items()
        }
        return jsonify({"settings": values})

    data = request.get_json() or {}
    updated = {}
    for key, value in data.items():
        if key not in allowed_keys:
            continue
        default, value_type, description = allowed_keys[key]
        AppSettings.set(key, value, value_type=value_type, category="ui", description=description)
        updated[key] = value

    if updated:
        AuditLog.log("site_controls_update", "AppSettings", "site_controls", updated)

    return jsonify({"message": "Settings updated", "updated": updated})


@app.route("/admin/audit-logs", methods=["GET"])
@login_required
def admin_audit_logs():
    """Admin: View all audit logs"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    action_filter = request.args.get("action")
    limit = request.args.get("limit", 100, type=int)
    offset = request.args.get("offset", 0, type=int)

    query = AuditLog.query

    if action_filter:
        query = query.filter_by(action=action_filter)

    total = query.count()
    logs = query.order_by(AuditLog.created_at.desc()).limit(limit).offset(offset).all()

    return jsonify(
        {
            "total": total,
            "limit": limit,
            "offset": offset,
            "logs": [
                {
                    "id": log.id,
                    "user_id": log.user_id,
                    "action": log.action,
                    "resource_type": log.resource_type,
                    "resource_id": log.resource_id,
                    "details": log.details,
                    "ip_address": log.ip_address,
                    "user_agent": log.user_agent,
                    "created_at": log.created_at.isoformat(),
                }
                for log in logs
            ],
        }
    )


@app.route("/admin/system-info", methods=["GET"])
@login_required
def admin_system_info():
    """Admin: Get system information"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    import psutil

    # Database size
    db_path = os.path.join(app.instance_path, "barberx_legal.db")
    db_size_mb = os.path.getsize(db_path) / (1024 * 1024) if os.path.exists(db_path) else 0

    # Upload folder size
    upload_folder = app.config["UPLOAD_FOLDER"]
    upload_size_mb = 0
    if os.path.exists(upload_folder):
        for root, dirs, files in os.walk(upload_folder):
            upload_size_mb += sum(os.path.getsize(os.path.join(root, f)) for f in files)
        upload_size_mb = upload_size_mb / (1024 * 1024)

    # System metrics
    cpu_percent = psutil.cpu_percent(interval=1)
    memory = psutil.virtual_memory()
    disk = psutil.disk_usage("/")

    return jsonify(
        {
            "database_size_mb": round(db_size_mb, 2),
            "upload_storage_mb": round(upload_size_mb, 2),
            "upload_storage_gb": round(upload_size_mb / 1024, 2),
            "cpu_percent": cpu_percent,
            "memory_percent": memory.percent,
            "memory_used_gb": round(memory.used / (1024**3), 2),
            "memory_total_gb": round(memory.total / (1024**3), 2),
            "disk_percent": disk.percent,
            "disk_used_gb": round(disk.used / (1024**3), 2),
            "disk_total_gb": round(disk.total / (1024**3), 2),
            "python_version": sys.version,
            "flask_version": flask.__version__,
        }
    )


# ========================================
# INITIALIZE DATABASE
# ========================================

try:
    with app.app_context():
        db.create_all()
        app.logger.info("Database tables initialized")

        # Initialize backend optimization services (indexes, evidence processor)
        initialize_backend_services()

        # Admin user already created via create_admin.py
        # Use admin@barberx.info with the 33-char password from that script
except Exception as e:
    print(f"[CRITICAL] Failed to initialize app: {e}")
    import traceback

    traceback.print_exc()
    # Don't crash - let app start but log the error
    app.logger.error(f"Initialization error: {e}")


# ========================================
# ADMIN SETTINGS MANAGEMENT
# ========================================


@app.route("/admin/settings", methods=["GET"])
@login_required
def admin_get_settings():
    """Admin: Get all app settings"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    # Get all settings grouped by category
    settings = AppSettings.query.order_by(AppSettings.category, AppSettings.key).all()

    settings_by_category = {}
    for setting in settings:
        if setting.category not in settings_by_category:
            settings_by_category[setting.category] = []

        settings_by_category[setting.category].append(
            {
                "id": setting.id,
                "key": setting.key,
                "value": setting.value,
                "value_type": setting.value_type,
                "description": setting.description,
                "is_editable": setting.is_editable,
                "updated_at": setting.updated_at.isoformat() if setting.updated_at else None,
            }
        )

    return jsonify({"settings": settings_by_category, "total": len(settings)})


@app.route("/admin/settings/<int:setting_id>", methods=["PUT"])
@login_required
def admin_update_setting(setting_id):
    """Admin: Update a setting"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    setting = AppSettings.query.get_or_404(setting_id)

    if not setting.is_editable:
        return jsonify({"error": "This setting is not editable"}), 400

    data = request.get_json()

    if "value" in data:
        setting.value = str(data["value"])
        setting.updated_at = datetime.utcnow()
        setting.updated_by = current_user.id

        db.session.commit()

        AuditLog.log(
            "setting_update",
            "AppSettings",
            str(setting_id),
            {"key": setting.key, "new_value": data["value"]},
        )

        return jsonify(
            {
                "message": "Setting updated successfully",
                "setting": {
                    "id": setting.id,
                    "key": setting.key,
                    "value": setting.value,
                    "updated_at": setting.updated_at.isoformat(),
                },
            }
        )

    return jsonify({"error": "No value provided"}), 400


@app.route("/admin/settings", methods=["POST"])
@login_required
def admin_create_setting():
    """Admin: Create a new setting"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    data = request.get_json()

    required_fields = ["key", "value", "category"]
    if not all(field in data for field in required_fields):
        return jsonify({"error": "Missing required fields"}), 400

    # Check if key already exists
    existing = AppSettings.query.filter_by(key=data["key"]).first()
    if existing:
        return jsonify({"error": "Setting key already exists"}), 400

    setting = AppSettings(
        key=data["key"],
        value=str(data["value"]),
        value_type=data.get("value_type", "string"),
        category=data["category"],
        description=data.get("description", ""),
        is_editable=data.get("is_editable", True),
        updated_by=current_user.id,
    )

    db.session.add(setting)
    db.session.commit()

    AuditLog.log(
        "setting_create",
        "AppSettings",
        str(setting.id),
        {"key": setting.key, "value": data["value"]},
    )

    return (
        jsonify(
            {
                "message": "Setting created successfully",
                "setting": {
                    "id": setting.id,
                    "key": setting.key,
                    "value": setting.value,
                    "category": setting.category,
                },
            }
        ),
        201,
    )


@app.route("/admin/settings/<int:setting_id>", methods=["DELETE"])
@login_required
def admin_delete_setting(setting_id):
    """Admin: Delete a setting"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    setting = AppSettings.query.get_or_404(setting_id)

    if not setting.is_editable:
        return jsonify({"error": "This setting cannot be deleted"}), 400

    key = setting.key
    db.session.delete(setting)
    db.session.commit()

    AuditLog.log("setting_delete", "AppSettings", str(setting_id), {"key": key})

    return jsonify({"message": "Setting deleted successfully"})


@app.route("/admin/settings/initialize", methods=["POST"])
@login_required
def admin_initialize_settings():
    """Admin: Initialize default settings"""
    if current_user.role != "admin":
        return jsonify({"error": "Admin access required"}), 403

    # Define default settings
    default_settings = [
        # General Settings
        {
            "key": "app_name",
            "value": "BarberX Legal Tech",
            "value_type": "string",
            "category": "general",
            "description": "Application name displayed throughout the platform",
        },
        {
            "key": "app_tagline",
            "value": "Professional BWC Forensic Analysis",
            "value_type": "string",
            "category": "general",
            "description": "Tagline shown on homepage and login",
        },
        {
            "key": "maintenance_mode",
            "value": "false",
            "value_type": "bool",
            "category": "general",
            "description": "Enable to put site in maintenance mode",
        },
        {
            "key": "allow_registrations",
            "value": "true",
            "value_type": "bool",
            "category": "general",
            "description": "Allow new user registrations",
        },
        {
            "key": "contact_email",
            "value": "support@barberx.info",
            "value_type": "string",
            "category": "general",
            "description": "Contact email for support",
        },
        # Security Settings
        {
            "key": "session_timeout_minutes",
            "value": "60",
            "value_type": "int",
            "category": "security",
            "description": "User session timeout in minutes",
        },
        {
            "key": "password_min_length",
            "value": "8",
            "value_type": "int",
            "category": "security",
            "description": "Minimum password length",
        },
        {
            "key": "require_email_verification",
            "value": "true",
            "value_type": "bool",
            "category": "security",
            "description": "Require email verification for new accounts",
        },
        {
            "key": "max_login_attempts",
            "value": "5",
            "value_type": "int",
            "category": "security",
            "description": "Max failed login attempts before lockout",
        },
        {
            "key": "enable_2fa",
            "value": "false",
            "value_type": "bool",
            "category": "security",
            "description": "Enable two-factor authentication",
        },
        # Feature Flags
        {
            "key": "enable_api",
            "value": "true",
            "value_type": "bool",
            "category": "features",
            "description": "Enable API access for Pro/Enterprise users",
        },
        {
            "key": "enable_analytics",
            "value": "true",
            "value_type": "bool",
            "category": "features",
            "description": "Enable analytics dashboard",
        },
        {
            "key": "enable_export",
            "value": "true",
            "value_type": "bool",
            "category": "features",
            "description": "Enable data export functionality",
        },
        {
            "key": "enable_webhooks",
            "value": "false",
            "value_type": "bool",
            "category": "features",
            "description": "Enable webhook notifications",
        },
        # Tier Limits
        {
            "key": "free_tier_analyses",
            "value": "5",
            "value_type": "int",
            "category": "limits",
            "description": "Max analyses per month for free tier",
        },
        {
            "key": "free_tier_storage_mb",
            "value": "500",
            "value_type": "int",
            "category": "limits",
            "description": "Max storage in MB for free tier",
        },
        {
            "key": "pro_tier_analyses",
            "value": "100",
            "value_type": "int",
            "category": "limits",
            "description": "Max analyses per month for professional tier",
        },
        {
            "key": "pro_tier_storage_mb",
            "value": "2048",
            "value_type": "int",
            "category": "limits",
            "description": "Max storage in MB for professional tier",
        },
        {
            "key": "max_file_size_mb",
            "value": "500",
            "value_type": "int",
            "category": "limits",
            "description": "Maximum file upload size in MB",
        },
        # Email Settings
        {
            "key": "smtp_enabled",
            "value": "false",
            "value_type": "bool",
            "category": "email",
            "description": "Enable SMTP email sending",
        },
        {
            "key": "smtp_host",
            "value": "",
            "value_type": "string",
            "category": "email",
            "description": "SMTP server hostname",
        },
        {
            "key": "smtp_port",
            "value": "587",
            "value_type": "int",
            "category": "email",
            "description": "SMTP server port",
        },
        {
            "key": "smtp_username",
            "value": "",
            "value_type": "string",
            "category": "email",
            "description": "SMTP username",
        },
        {
            "key": "from_email",
            "value": "noreply@barberx.info",
            "value_type": "string",
            "category": "email",
            "description": "From email address",
        },
        # Branding
        {
            "key": "primary_color",
            "value": "#3b82f6",
            "value_type": "string",
            "category": "branding",
            "description": "Primary brand color (hex)",
        },
        {
            "key": "secondary_color",
            "value": "#8b5cf6",
            "value_type": "string",
            "category": "branding",
            "description": "Secondary brand color (hex)",
        },
        {
            "key": "logo_url",
            "value": "/assets/images/logo.png",
            "value_type": "string",
            "category": "branding",
            "description": "URL to application logo",
        },
        {
            "key": "favicon_url",
            "value": "/assets/images/favicon.ico",
            "value_type": "string",
            "category": "branding",
            "description": "URL to favicon",
        },
    ]

    created = 0
    skipped = 0

    for setting_data in default_settings:
        existing = AppSettings.query.filter_by(key=setting_data["key"]).first()
        if not existing:
            setting = AppSettings(
                key=setting_data["key"],
                value=setting_data["value"],
                value_type=setting_data["value_type"],
                category=setting_data["category"],
                description=setting_data["description"],
                is_editable=True,
                updated_by=current_user.id,
            )
            db.session.add(setting)
            created += 1
        else:
            skipped += 1

    db.session.commit()

    AuditLog.log(
        "settings_initialize", "AppSettings", None, {"created": created, "skipped": skipped}
    )

    return jsonify(
        {
            "message": "Settings initialized successfully",
            "created": created,
            "skipped": skipped,
            "total": created + skipped,
        }
    )


# ========================================
# HEALTH CHECK ENDPOINT
# ========================================

# ========================================
# AI CHAT ENDPOINT (for chat widget)
# ========================================
try:
    import openai

    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    app.logger.warning("OpenAI not available - AI chat features disabled")


from flask_login import login_required

# Subscription & tier gating imports


@app.route("/api/upload/pdf/secure", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("pdf_documents_per_month", increment=1)
def upload_pdf_secure():
    try:
        from pypdf import PdfReader
    except ImportError:
        return jsonify({"error": "PDF processing not available"}), 503

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF files allowed"}), 400
    save_path = app.config["UPLOAD_FOLDER"].parent / "pdfs" / secure_filename(file.filename)
    file.save(save_path)
    # Extract text for chat context
    with open(save_path, "rb") as f:
        reader = PdfReader(f)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return jsonify(
        {"message": "PDF uploaded", "filename": str(save_path.name), "text": text[:10000]}
    )


@app.route("/api/upload/video", methods=["POST"])
@login_required
@require_tier(TierLevel.STARTER)
@check_usage_limit("bwc_videos_per_month", increment=1)
def upload_video():
    """Upload video and generate mock transcription"""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]

    if not file.filename.lower().endswith((".mp4", ".mov", ".avi", ".webm")):
        return jsonify({"error": "Only video files allowed (.mp4, .mov, .avi, .webm)"}), 400

    # Secure filename
    filename = secure_filename(file.filename)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{current_user.id}_{timestamp}_{filename}"

    # Save to BWC videos folder
    upload_dir = Path(app.config.get("UPLOAD_FOLDER", "./uploads/bwc_videos"))
    upload_dir.mkdir(parents=True, exist_ok=True)
    filepath = upload_dir / unique_filename
    file.save(filepath)

    # Calculate file hash
    file_hash = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            file_hash.update(chunk)
    file_hash_hex = file_hash.hexdigest()

    file_size = os.path.getsize(filepath)

    # Create analysis record
    analysis = Analysis(
        user_id=current_user.id,
        filename=filename,
        file_hash=file_hash_hex,
        file_size=file_size,
        file_path=str(filepath),
        status="uploaded",
    )
    analysis.generate_id()

    # Get optional metadata from form
    analysis.case_number = request.form.get("case_number", "")
    analysis.evidence_number = request.form.get("evidence_number", "")
    analysis.acquired_by = request.form.get(
        "acquired_by", current_user.full_name or current_user.email
    )
    analysis.source = request.form.get("source", "Web Upload")

    db.session.add(analysis)
    db.session.commit()

    # Start enhanced analysis in background
    try:
        from enhanced_analysis_tools import advanced_analyzer

        def generate_mock_analysis():
            try:
                analysis.status = "analyzing"
                analysis.current_step = "Generating enhanced transcription..."
                analysis.progress = 10
                db.session.commit()

                # Generate enhanced analysis
                import time

                time.sleep(2)  # Simulate processing

                analysis.progress = 30
                analysis.current_step = "Analyzing voice stress patterns..."
                db.session.commit()

                time.sleep(1)

                analysis.progress = 50
                analysis.current_step = "Detecting scenes and objects..."
                db.session.commit()

                time.sleep(1)

                # Generate complete enhanced analysis
                mock_report = advanced_analyzer.generate_complete_analysis(
                    str(filepath),
                    case_number=analysis.case_number,
                    evidence_number=analysis.evidence_number,
                )

                analysis.progress = 80
                analysis.current_step = "Generating reports and visualizations..."
                db.session.commit()

                # Save report to file
                output_dir = Path(app.config.get("ANALYSIS_FOLDER", "./bwc_analysis")) / analysis.id
                output_dir.mkdir(parents=True, exist_ok=True)

                # Save JSON report
                import json

                json_path = output_dir / "report.json"
                with open(json_path, "w") as f:
                    json.dump(mock_report, f, indent=2)

                # Save TXT report
                txt_path = output_dir / "report.txt"
                with open(txt_path, "w") as f:
                    f.write("BWC FORENSIC ANALYSIS REPORT\n")
                    f.write("=" * 60 + "\n\n")
                    f.write(f"Case Number: {mock_report['metadata']['case_number']}\n")
                    f.write(f"Evidence Number: {mock_report['metadata']['evidence_number']}\n")
                    f.write(f"Duration: {mock_report['metadata']['duration']} seconds\n")
                    f.write(f"Speakers: {len(mock_report['transcript']['speakers'])}\n")
                    f.write(f"Segments: {len(mock_report['transcript']['segments'])}\n")
                    f.write(f"Discrepancies: {len(mock_report['discrepancies']['items'])}\n\n")
                    f.write("TRANSCRIPT\n")
                    f.write("-" * 60 + "\n\n")
                    for segment in mock_report["transcript"]["segments"]:
                        f.write(
                            f"[{segment['start_time']:.2f}s] {segment['speaker_label']}: {segment['text']}\n"
                        )

                # Update analysis record
                analysis.status = "completed"
                analysis.progress = 100
                analysis.current_step = "Analysis complete"
                analysis.completed_at = datetime.utcnow()
                analysis.duration = mock_report["metadata"]["duration"]
                analysis.total_speakers = len(mock_report["transcript"]["speakers"])
                analysis.total_segments = len(mock_report["transcript"]["segments"])
                analysis.total_discrepancies = len(mock_report["discrepancies"]["items"])
                analysis.critical_discrepancies = mock_report["discrepancies"]["by_severity"][
                    "high"
                ]
                analysis.report_json_path = str(json_path)
                analysis.report_txt_path = str(txt_path)
                analysis.report_md_path = str(output_dir / "report.md")

                db.session.commit()

                app.logger.info(f"Mock analysis completed for: {analysis.id}")

            except Exception as e:
                analysis.status = "failed"
                analysis.error_message = str(e)
                analysis.progress = 0
                db.session.commit()
                app.logger.error(f"Mock analysis failed: {e}")

        # Start background thread
        import threading

        thread = threading.Thread(target=generate_mock_analysis)
        thread.daemon = True
        thread.start()

        app.logger.info(f"Video uploaded and mock analysis started: {filename}")

    except ImportError:
        # If mock_analysis not available, just mark as uploaded
        app.logger.warning("Mock analysis generator not available")

    return jsonify(
        {
            "success": True,
            "message": "Video uploaded successfully. Analysis starting...",
            "upload_id": analysis.id,
            "filename": filename,
            "file_hash": file_hash_hex,
            "file_size": file_size,
            "status": "analyzing",
        }
    )


@app.route("/api/chat", methods=["POST"])
@login_required
def ai_chat():
    import time

    from usage_meter import SmartMeter, UsageQuota

    # Check quota and rate limit
    quota = UsageQuota.query.filter_by(user_id=current_user.id).first()
    if not quota:
        quota = SmartMeter.initialize_user_quota(current_user.id)

    # Rate limit check
    if not quota.check_rate_limit():
        return (
            jsonify(
                {
                    "error": "Rate limit exceeded",
                    "message": "Too many requests. Please wait a moment.",
                    "retry_after": 60,
                }
            ),
            429,
        )

    # AI request quota check
    has_quota, error_msg = quota.check_quota("ai_requests")
    if not has_quota:
        SmartMeter.track_event(
            event_type="chat_blocked",
            event_category="quota",
            status="denied",
            error_message=error_msg,
        )
        return (
            jsonify(
                {"error": "AI quota exceeded", "message": error_msg, "upgrade_url": "/pricing"}
            ),
            429,
        )

    quota.requests_this_minute += 1
    quota.last_request_timestamp = datetime.utcnow()
    db.session.commit()

    data = request.get_json()
    question = data.get("question", "").strip()
    context = data.get("context", "")
    user_api_key = data.get("api_key", "").strip()
    model = data.get("model", "gpt-4")

    if not question:
        return jsonify({"error": "No question provided"}), 400

    # Use user-provided API key if present, else server key
    api_key = user_api_key or os.getenv("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"answer": "[No OpenAI API key provided. Please enter your key.]"}), 400

    start_time = time.time()
    tokens_used = 0

    try:
        openai.api_key = api_key
        prompt = f"You are a legal tech assistant. Context: {context}. User question: {question}"
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful legal tech assistant for BWC and analysis reports.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=400,
            temperature=0.2,
        )
        answer = response.choices[0].message["content"].strip()
        tokens_used = response.get("usage", {}).get("total_tokens", 0)

        duration = time.time() - start_time

        # Estimate cost
        cost_per_1k = 0.03 if "gpt-4" in model else 0.0015
        estimated_cost = (tokens_used / 1000) * cost_per_1k

        # Track successful chat
        SmartMeter.track_event(
            event_type="chat_message",
            event_category="compute",
            resource_name=model,
            tokens_input=len(question.split()),
            tokens_output=tokens_used,
            duration_seconds=duration,
            cost_usd=estimated_cost,
            status="success",
        )

        # Update quotas
        quota.increment_quota("ai_requests")
        quota.increment_quota("ai_tokens", tokens_used)
        quota.total_cost_usd += estimated_cost
        db.session.commit()

    except Exception as e:
        duration = time.time() - start_time
        answer = f"[AI unavailable: {e}]"

        # Track error
        SmartMeter.track_event(
            event_type="chat_message",
            event_category="compute",
            resource_name=model,
            duration_seconds=duration,
            status="error",
            error_message=str(e),
        )

    return jsonify({"answer": answer, "tokens_used": tokens_used})


@app.route("/healthz")
def health_check_status():
    """Health check endpoint for monitoring"""
    return jsonify(
        {
            "status": "healthy",
            "timestamp": datetime.utcnow().isoformat(),
            "version": "1.0.0",
            "database": "connected" if db.engine else "disconnected",
        }
    )


# ========================================
# RUN APPLICATION


@app.route("/dashboard/usage")
@login_required
def usage_dashboard():
    """Usage dashboard showing subscription and limits"""
    from tier_gating import TierGate

    usage_stats = TierGate.get_usage_stats(current_user)

    return render_template("usage_dashboard.html", usage_stats=usage_stats)


# ========================================


# ============================================================================
# FREE TIER ROUTES
# ============================================================================


@app.route("/free-dashboard")
@login_required
def free_dashboard():
    """FREE tier dashboard with demo cases and educational resources"""
    from flask_login import current_user

    from models_auth import TierLevel

    # Redirect non-FREE users to regular dashboard
    if current_user.tier != TierLevel.FREE:
        return redirect(url_for("dashboard"))

    # Get demo cases
    demo_cases = get_demo_cases()

    # Get educational resources
    educational_resources = get_all_educational_resources()

    # Get upload status
    upload_status = OneTimeUploadManager.get_upload_status(current_user)

    # Get data retention status
    data_status = get_user_data_status(current_user)

    return render_template(
        "free_tier_dashboard.html",
        demo_cases=demo_cases,
        educational_resources=educational_resources,
        upload_status=upload_status,
        data_status=data_status,
    )


@app.route("/cases/<case_id>")
@login_required
def view_case(case_id):
    """View case details (handles both demo and real cases)"""
    from flask_login import current_user

    # Check if it's a demo case
    if is_demo_case(case_id):
        demo_case = get_demo_case_by_id(case_id)
        if not demo_case:
            flash("Demo case not found", "error")
            return redirect(url_for("free_dashboard"))

        return render_template("demo_case_detail.html", case=demo_case)

    # Regular case logic here
    # case = Case.query.get_or_404(case_id)
    # return render_template('case_detail.html', case=case)

    flash("Case viewing coming soon", "info")
    return redirect(url_for("free_dashboard"))


@app.route("/education")
@app.route("/education/<category>")
@login_required
def education_center(category=None):
    """Educational resources center"""
    resources = get_all_educational_resources()

    return render_template(
        "education_center.html",
        resources=resources,
        categories=CATEGORIES,
        selected_category=category,
    )


@app.route("/education/resource/<resource_id>")
@login_required
def view_resource(resource_id):
    """View specific educational resource"""
    resource = get_resource_by_id(resource_id)

    if not resource:
        flash("Resource not found", "error")
        return redirect(url_for("education_center"))

    return render_template("resource_detail.html", resource=resource)


@app.route("/api/upload-status")
@login_required
def get_upload_status():
    """API endpoint for upload status"""
    from flask import jsonify
    from flask_login import current_user

    status = OneTimeUploadManager.get_upload_status(current_user)
    return jsonify(status)


@app.route("/api/data-retention-status")
@login_required
def get_data_retention_status():
    """API endpoint for data retention status"""
    from flask import jsonify
    from flask_login import current_user

    status = get_user_data_status(current_user)
    return jsonify(status)


# ============================================================================
# END FREE TIER ROUTES
# ============================================================================


if __name__ == "__main__":
    # Get port from environment variable (for cloud deployments)
    port = int(os.getenv("PORT", 5000))
    debug = os.getenv("FLASK_ENV") != "production"

    # Initialize backend optimization services
    initialize_backend_services()

    print(
        f"""
    ================================================================

           BarberX Legal Technologies
           Professional BWC Forensic Analysis Platform

    ================================================================

    Web Application: http://localhost:{port}
    Admin Login: admin@barberx.info
    Database: SQLite (Development)
    Logs: ./logs/barberx.log

    BACKEND OPTIMIZATION:
    [+] Database connection pooling (30 concurrent users)
    [+] Intelligent caching (99% faster on hits)
    [+] API rate limiting (tiered by subscription)
    [+] Performance monitoring
    [+] Unified evidence processing pipeline

    ENHANCED FEATURES:
    [+] Video analysis with synchronized transcript
    [+] Voice stress analysis and emotion detection
    [+] Automatic scene detection
    [+] Audio/video quality assessment
    [+] Legal compliance checking
    [+] Interactive timeline and waveform
    [+] Terminal interface (Ctrl+`)
    [+] AI chat assistant
    [+] Annotation system
    [+] Multi-format report export

    CORE FEATURES:
    [+] Multi-user authentication
    [+] Role-based access control
    [+] Subscription tiers (Free, Professional, Enterprise)
    [+] API key management
    [+] Audit logging

    Ready for production deployment!
    Press Ctrl+C to stop the server.
    """.format(
            port=port
        )
    )

    app.run(host="0.0.0.0", port=port, debug=debug)


# ========================================
# RUN APPLICATION
# ========================================


# Stripe Pricing Table Routes
@app.route("/pricing-stripe")
def pricing_stripe():
    """Stripe embedded pricing table"""
    return render_template("pricing-stripe-embed.html")


@app.route("/pricing")
def pricing():
    """Main pricing page with custom cards"""
    # Option 1: Use custom cards (pricing-5tier.html)
    return render_template("pricing.html")

    # Option 2: Use Stripe embed (uncomment to switch)
    # return render_template('pricing-stripe-embed.html')

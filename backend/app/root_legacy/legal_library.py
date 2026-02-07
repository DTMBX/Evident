from typing import Optional
# Copyright © 2024–2026 Faith Frontier Ecclesiastical Trust. All rights reserved.
# PROPRIETARY — See LICENSE.

"""
Legal Reference Library - Comprehensive case law and statute database

This module provides:
- Storage for legal documents (case law, statutes, regulations)
- Full-text search across legal materials
- Citation parsing and linking
- Web scraping from legal databases (CourtListener, Justia, etc.)
- Integration with ChatGPT for intelligent legal research
- Document ingestion from PDF/TXT/DOCX files

Supports:
- Supreme Court opinions
- Federal Circuit Courts
- State courts
- Statutes (federal and state)
- Regulations (CFR, state regs)
- Law review articles
- User-uploaded legal documents
"""

import json
import os
import re
from datetime import datetime

import docx
import pypdf as PyPDF2  # Migrated from PyPDF2 (deprecated), aliased for compatibility
import requests
from bs4 import BeautifulSoup
from sqlalchemy import Boolean, Column, DateTime, ForeignKey, Integer, String, Text
from sqlalchemy.orm import relationship

# Database Models
from .models_auth import db  # Use existing database


class LegalDocument(db.Model):
    """Base model for all legal documents"""

    __tablename__ = "legal_documents"

    id = Column(Integer, primary_key=True)

    # Document identification
    doc_type = Column(
        String(50), nullable=False
    )  # 'case', 'statute', 'regulation', 'article', 'user_upload'
    title = Column(String(500), nullable=False)
    citation = Column(String(200), unique=True)  # e.g., "Brown v. Board, 347 U.S. 483 (1954)"

    # Court/Source information
    court = Column(String(200))  # e.g., "U.S. Supreme Court"
    jurisdiction = Column(String(100))  # e.g., "Federal", "California"
    decision_date = Column(DateTime)

    # Content
    full_text = Column(Text, nullable=False)
    summary = Column(Text)  # Brief summary/headnote

    # Metadata
    judges = Column(String(500))  # Comma-separated
    case_number = Column(String(100))  # Docket number
    url = Column(String(500))  # Source URL

    # Search and categorization
    topics = Column(Text)  # JSON array of topics: ["4th Amendment", "Search and Seizure"]
    legal_issues = Column(Text)  # JSON array: ["Warrantless search", "Probable cause"]

    # User/case association
    user_id = Column(Integer)  # No foreign key - just reference
    case_id = Column(Integer)  # No foreign key - optional case association

    # Upload info
    uploaded_by = Column(Integer)  # No foreign key - just reference
    uploaded_at = Column(DateTime, default=datetime.utcnow)
    source = Column(String(100))  # 'courtlistener', 'justia', 'user_upload', 'web_scrape'

    # Status
    verified = Column(Boolean, default=False)  # Verified by user/admin
    public = Column(Boolean, default=True)  # Available to all users vs. private

    # Relationships
    citations_made = relationship(
        "Citation", foreign_keys="Citation.citing_doc_id", backref="citing_document"
    )
    citations_received = relationship(
        "Citation", foreign_keys="Citation.cited_doc_id", backref="cited_document"
    )
    annotations = relationship(
        "DocumentAnnotation", backref="document", cascade="all, delete-orphan"
    )

    def __repr__(self):
        return f"<LegalDocument {self.citation}: {self.title}>"


class Citation(db.Model):
    """Links between documents (case A cites case B)"""

    __tablename__ = "citations"

    id = Column(Integer, primary_key=True)
    citing_doc_id = Column(Integer, ForeignKey("legal_documents.id"), nullable=False)
    cited_doc_id = Column(Integer, ForeignKey("legal_documents.id"), nullable=False)
    context = Column(Text)  # Text surrounding the citation
    citation_type = Column(String(50))  # 'positive', 'negative', 'neutral', 'distinguished'
    created_at = Column(DateTime, default=datetime.utcnow)


class DocumentAnnotation(db.Model):
    """User annotations on legal documents"""

    __tablename__ = "document_annotations"

    id = Column(Integer, primary_key=True)
    document_id = Column(Integer, ForeignKey("legal_documents.id"), nullable=False)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)

    text_selection = Column(Text)  # The text that was highlighted
    annotation = Column(Text)  # User's note
    tags = Column(String(500))  # Comma-separated tags

    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class LegalTopic(db.Model):
    """Taxonomy of legal topics"""

    __tablename__ = "legal_topics"

    id = Column(Integer, primary_key=True)
    name = Column(String(200), unique=True, nullable=False)
    parent_id = Column(Integer, ForeignKey("legal_topics.id"))  # For hierarchical topics
    description = Column(Text)

    # Relationships
    children = relationship("LegalTopic", backref=db.backref("parent", remote_side=[id]))


# Service Classes


class LegalLibraryService:
    """Main service for legal library operations"""

    def __init__(self):
        self.courtlistener_api = "https://www.courtlistener.com/api/rest/v4/"
        self.justia_api = "https://api.justia.com/"

    def search_library(
        self,
        query: str,
Optional[doc_type: str] = None,
Optional[court: str] = None,
Optional[jurisdiction: str] = None,
Optional[date_from: datetime] = None,
Optional[date_to: datetime] = None,
        limit: int = 50,
    ) -> list[LegalDocument]:
        """
        Full-text search across legal library

        Examples:
        - search_library("fourth amendment search warrant")
        - search_library("miranda rights", court="U.S. Supreme Court")
        - search_library("excessive force", jurisdiction="California")
        """

        # Build query
        query_obj = LegalDocument.query

        # Full-text search (using LIKE for SQLite, would use FTS for production)
        query_obj = query_obj.filter(
            db.or_(
                LegalDocument.title.contains(query),
                LegalDocument.full_text.contains(query),
                LegalDocument.summary.contains(query),
            )
        )

        # Filters
        if doc_type:
            query_obj = query_obj.filter(LegalDocument.doc_type == doc_type)
        if court:
            query_obj = query_obj.filter(LegalDocument.court.contains(court))
        if jurisdiction:
            query_obj = query_obj.filter(LegalDocument.jurisdiction == jurisdiction)
        if date_from:
            query_obj = query_obj.filter(LegalDocument.decision_date >= date_from)
        if date_to:
            query_obj = query_obj.filter(LegalDocument.decision_date <= date_to)

        # Execute
        results = query_obj.order_by(LegalDocument.decision_date.desc()).limit(limit).all()

        return results

Optional[def ingest_from_courtlistener(self, citation: str) -> LegalDocument]:
        """
        Fetch case from CourtListener API by citation

        Example: ingest_from_courtlistener("347 U.S. 483")

        Requires COURTLISTENER_API_KEY environment variable
        Get your free API key at: https://www.courtlistener.com/api/rest-info/
        """

        # CourtListener API endpoint
        url = f"{self.courtlistener_api}search/"
        params = {"citation": citation, "format": "json"}

        # Add API key from environment if available
        headers = {}
        api_key = os.getenv("COURTLISTENER_API_KEY")
        if api_key:
            headers["Authorization"] = f"Token {api_key}"

        try:
            response = requests.get(url, params=params, headers=headers)
            response.raise_for_status()

            data = response.json()

            if not data.get("results"):
                return None

            # Parse first result
            result = data["results"][0]

            doc = LegalDocument(
                doc_type="case",
                title=result.get("caseName", ""),
                citation=citation,
                court=result.get("court", ""),
                jurisdiction=result.get("jurisdiction", "Federal"),
                decision_date=(
                    datetime.fromisoformat(result.get("dateFiled"))
                    if result.get("dateFiled")
                    else None
                ),
                full_text=result.get("plain_text", ""),
                summary=result.get("snippet", ""),
                judges=", ".join(result.get("judges", [])),
                case_number=result.get("docketNumber", ""),
                url=result.get("absolute_url", ""),
                source="courtlistener",
                verified=True,
                public=True,
            )

            # Extract topics and legal issues (would use NLP in production)
            doc.topics = json.dumps(self._extract_topics(doc.full_text))
            doc.legal_issues = json.dumps(self._extract_legal_issues(doc.full_text))

            # Save to database
            db.session.add(doc)
            db.session.commit()

            return doc

        except Exception as e:
            print(f"Error fetching from CourtListener: {e}")
            return None

    def ingest_from_file(
        self,
        file_path: str,
        title: str,
        doc_type: str = "user_upload",
Optional[user_id: int] = None,
Optional[case_id: int] = None,
Optional[metadata: dict] = None,
Optional[) -> LegalDocument]:
        """
        Ingest legal document from uploaded file (PDF, TXT, DOCX)

        Example:
        ingest_from_file(
            "uploads/miranda_v_arizona.pdf",
            "Miranda v. Arizona",
            doc_type='case',
            user_id=1
        )
        """

        # Extract text from file
        text = self._extract_text_from_file(file_path)

        if not text:
            return None

        # Parse citation if present
        citation = self._parse_citation(text) or metadata.get("citation") if metadata else None

        # Create document
        doc = LegalDocument(
            doc_type=doc_type,
            title=title,
            citation=citation,
            full_text=text,
            uploaded_by=user_id,
            user_id=user_id,
            case_id=case_id,
            source="user_upload",
            public=metadata.get("public", False) if metadata else False,
        )

        # Apply metadata
        if metadata:
            if metadata.get("court"):
                doc.court = metadata["court"]
            if metadata.get("jurisdiction"):
                doc.jurisdiction = metadata["jurisdiction"]
            if metadata.get("decision_date"):
                doc.decision_date = metadata["decision_date"]
            if metadata.get("summary"):
                doc.summary = metadata["summary"]

        # Extract topics and issues
        doc.topics = json.dumps(self._extract_topics(text))
        doc.legal_issues = json.dumps(self._extract_legal_issues(text))

        # Save
        db.session.add(doc)
        db.session.commit()

        # Extract and link citations
        self._extract_and_link_citations(doc)

        return doc

    def search_web_for_case(self, query: str, source: str = "justia") -> list[dict]:
        """
        Search web for legal cases (Justia, Google Scholar, etc.)

        Returns list of results that can be ingested
        """

        if source == "justia":
            return self._search_justia(query)
        elif source == "scholar":
            return self._search_google_scholar(query)
        else:
            return []

    def get_related_cases(self, doc_id: int, limit: int = 10) -> list[LegalDocument]:
        """
        Find related cases based on:
        - Similar topics
        - Cited by / cites
        - Same jurisdiction
        - Same court
        """

        doc = LegalDocument.query.get(doc_id)
        if not doc:
            return []

        # Get topics
        topics = json.loads(doc.topics) if doc.topics else []

        # Find docs with overlapping topics
        related = (
            LegalDocument.query.filter(
                LegalDocument.id != doc_id,
                LegalDocument.topics.contains(topics[0]) if topics else True,
            )
            .limit(limit)
            .all()
        )

        return related

    def annotate_document(
        self,
        doc_id: int,
        user_id: int,
        text_selection: str,
        annotation: str,
Optional[tags: list[str]] = None,
    ) -> DocumentAnnotation:
        """Add user annotation to document"""

        annot = DocumentAnnotation(
            document_id=doc_id,
            user_id=user_id,
            text_selection=text_selection,
            annotation=annotation,
            tags=", ".join(tags) if tags else None,
        )

        db.session.add(annot)
        db.session.commit()

        return annot

    # Private helper methods

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from PDF, TXT, or DOCX"""

        ext = file_path.lower().split(".")[-1]

        try:
            if ext == "pdf":
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text

            elif ext == "docx":
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])

            elif ext == "txt":
                with open(file_path, encoding="utf-8") as f:
                    return f.read()

            else:
                return None

        except Exception as e:
            print(f"Error extracting text: {e}")
            return None

Optional[def _parse_citation(self, text: str) -> str]:
        """Extract legal citation from text"""

        # Common citation patterns
        patterns = [
            r"\d+\s+U\.S\.\s+\d+",  # U.S. Reports: "347 U.S. 483"
            r"\d+\s+S\.\s?Ct\.\s+\d+",  # Supreme Court Reporter
            r"\d+\s+F\.\d+d\s+\d+",  # Federal Reporter
            r"\d+\s+Cal\.\s?\d+d\s+\d+",  # California Reports
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)

        return None

    def _extract_topics(self, text: str) -> list[str]:
        """Extract legal topics from text (simplified)"""

        topics = []

        # Keyword matching (would use NLP in production)
        topic_keywords = {
            "4th Amendment": [
                "fourth amendment",
                "search and seizure",
                "warrant",
                "probable cause",
            ],
            "5th Amendment": ["fifth amendment", "self-incrimination", "miranda"],
            "Civil Rights": ["civil rights", "1983", "constitutional violation"],
            "Excessive Force": ["excessive force", "use of force", "police brutality"],
            "Due Process": ["due process", "procedural", "substantive"],
        }

        text_lower = text.lower()

        for topic, keywords in topic_keywords.items():
            if any(kw in text_lower for kw in keywords):
                topics.append(topic)

        return topics[:5]  # Limit to top 5

    def _extract_legal_issues(self, text: str) -> list[str]:
        """Extract specific legal issues (simplified)"""

        issues = []

        # Pattern matching for common legal issues
        issue_patterns = {
            "Warrantless search": r"warrantless\s+search",
            "Qualified immunity": r"qualified\s+immunity",
            "Probable cause": r"probable\s+cause",
            "Miranda warnings": r"miranda\s+(warning|rights)",
        }

        for issue, pattern in issue_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                issues.append(issue)

        return issues

    def _extract_and_link_citations(self, doc: LegalDocument):
        """Extract citations from document and link to cited cases"""

        # Find all citations in text
        citation_pattern = r"(\d+\s+U\.S\.\s+\d+|\d+\s+F\.\d+d\s+\d+)"
        matches = re.finditer(citation_pattern, doc.full_text)

        for match in matches:
            cited_citation = match.group(1)

            # Find cited document in database
            cited_doc = LegalDocument.query.filter_by(citation=cited_citation).first()

            if cited_doc:
                # Create citation link
                citation = Citation(
                    citing_doc_id=doc.id,
                    cited_doc_id=cited_doc.id,
                    context=doc.full_text[max(0, match.start() - 100) : match.end() + 100],
                )
                db.session.add(citation)

        db.session.commit()

    def _search_justia(self, query: str) -> list[dict]:
        """Search Justia for cases"""

        # Justia search URL (web scraping since API is limited)
        url = "https://law.justia.com/search"
        params = {"q": query}

        try:
            response = requests.get(url, params=params)
            soup = BeautifulSoup(response.content, "html.parser")

            results = []
            for result in soup.find_all("div", class_="result")[:10]:
                title_elem = result.find("a")
                if title_elem:
                    results.append(
                        {
                            "title": title_elem.text.strip(),
                            "url": title_elem.get("href"),
                            "snippet": result.find("p").text.strip() if result.find("p") else "",
                        }
                    )

            return results

        except Exception as e:
            print(f"Error searching Justia: {e}")
            return []

    def _search_google_scholar(self, query: str) -> list[dict]:
        """Search Google Scholar for legal cases"""
        # Implementation would use Google Scholar API or web scraping
        # Placeholder for now
        return []


class CitationParser:
    """Parse and validate legal citations"""

    @staticmethod
Optional[def parse(citation_text: str) -> dict]:
        """
        Parse citation into components

        Example: "Brown v. Board of Education, 347 U.S. 483 (1954)"
        Returns: {
            'case_name': 'Brown v. Board of Education',
            'volume': '347',
            'reporter': 'U.S.',
            'page': '483',
            'year': '1954',
            'full_citation': '347 U.S. 483 (1954)'
        }
        """

        # Regex for common citation formats
        pattern = r"(.+?),?\s+(\d+)\s+([A-Za-z\.\s]+?)\s+(\d+)\s+\((\d{4})\)"

        match = re.search(pattern, citation_text)
        if match:
            return {
                "case_name": match.group(1).strip(),
                "volume": match.group(2),
                "reporter": match.group(3).strip(),
                "page": match.group(4),
                "year": match.group(5),
                "full_citation": f"{match.group(2)} {match.group(3).strip()} {match.group(4)} ({match.group(5)})",
            }

        return None

    @staticmethod
    def standardize(citation: str) -> str:
        """Standardize citation format"""
        parsed = CitationParser.parse(citation)
        if parsed:
            return parsed["full_citation"]
        return citation

    @staticmethod
    def is_valid(citation: str) -> bool:
        """Validate citation format"""
        return CitationParser.parse(citation) is not None